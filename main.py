######################################
####import library
from dotenv import load_dotenv
from elevate import elevate
import zipfile
import secrets
import sys, traceback
import mysql.connector
import sys,os,re
import json
from PIL import Image
from datetime import datetime
from PyQt5 import QtCore,QtGui
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.uic import loadUiType
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QPushButton, QTableWidget,QSizeGrip,QLineEdit,QAction,
    QTableWidgetItem, QDialog, QHBoxLayout, QFileDialog,QMessageBox, QMainWindow,QCompleter,QLabel
)
from PyQt5.QtCore import QDate,Qt,QSettings,Qt
from PyQt5.QtPrintSupport import QPrinter, QPrintDialog, QPrintPreviewDialog
from PyQt5.QtGui import QTextDocument,QPixmap, QFont,QPainter
from time import sleep
import pandas as pd
from openpyxl import Workbook
from openpyxl.drawing.image import Image
from openpyxl.styles import Alignment
import xlsxwriter
#####################################
###############For Documents####################
from docx import Document
from docx.shared import Inches
#################################
from capture_window import CaptureWindow
from qr_display import QRDisplay
from qr_scanner import QRScannerWindow
from qt_material import *
from qt_material import apply_stylesheet
###############################################
import icons_rc



def clean_png(path):
    img = Image.open(path)
    img.save(path, optimize=True)

################################################

class PrintPreviewDialog(QDialog):
    
    def __init__(self, table_widget, parent=None):
        super().__init__(parent)
        self.setWindowTitle("معاينة الطباعة")
        self.resize(800, 600)
        self.table_widget = table_widget

        # إنشاء المستند من الجدول
        self.document = QTextDocument()
        self.document.setHtml(self.generate_html_from_table())

        # الأزرار
        self.print_button = QPushButton("طباعة")
        self.save_button = QPushButton("حفظ كـ PDF")
        self.close_button = QPushButton("إغلاق")

        self.print_button.clicked.connect(self.print_document)
        self.save_button.clicked.connect(self.save_pdf)
        self.close_button.clicked.connect(self.close)

        # التخطيط
        layout = QVBoxLayout()
        layout.addWidget(self.table_widget)
        button_layout = QHBoxLayout()
        button_layout.addWidget(self.print_button)
        button_layout.addWidget(self.save_button)
        button_layout.addWidget(self.close_button)
        layout.addLayout(button_layout)
        self.setLayout(layout)

    def generate_html_from_table(self):
        # مسار صورة الشعار
        logo_path = "3kab.png"  # تأكد من وجود الصورة في نفس مجلد البرنامج

        html = f"""
    #     <div style='text-align: center;'>
    #         <img src="{logo_path}" width="100" height="100" />
    #         <h2>وزارةالمالية</h2>
    #         <h3> الهيئة العامة للضرائب والرسوم </h3>
    #         <h4>مديرية مالية حلب</h4>
    #         <h5>معاينة الجدول</h5>
    #     </div>
    #     <table border='1' cellspacing='0' cellpadding='4'>
    #     """

        # رؤوس الأعمدة
        html += "<tr>"
        for col in range(self.table_widget.columnCount()):
            header = self.table_widget.horizontalHeaderItem(col).text()
            html += f"<th>{header}</th>"
        html += "</tr>"

        # بيانات الجدول
        for row in range(self.table_widget.rowCount()):
            html += "<tr>"
            for col in range(self.table_widget.columnCount()):
                item = self.table_widget.item(row, col)
                html += f"<td>{item.text() if item else ''}</td>"
            html += "</tr>"

        html += "</table>"
        return html

    def print_document(self):
        printer = QPrinter(QPrinter.HighResolution)
        dialog = QPrintDialog(printer, self)
        if dialog.exec_() == QPrintDialog.Accepted:
            self.document.print_(printer)

    def save_pdf(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "حفظ كـ PDF", "", "PDF Files (*.pdf)")
        if file_path:
            printer = QPrinter(QPrinter.HighResolution)
            printer.setOutputFormat(QPrinter.PdfFormat)
            printer.setOutputFileName(file_path)
            self.document.print_(printer)

class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("البرنامج الرئيسي")
        self.resize(600, 400)

        self.print_preview_button = QPushButton("معاينة وطباعة")
        self.print_preview_button.clicked.connect(self.open_preview)

        layout = QVBoxLayout()
        layout.addWidget(self.table)
        layout.addWidget(self.print_preview_button)
        self.setLayout(layout)

    def open_preview(self):
        # إنشاء نسخة من الجدول لعرضها في نافذة المعاينة
        preview_table = QTableWidget(self.table.rowCount(), self.table.columnCount())
        preview_table.setHorizontalHeaderLabels(
            [self.table.horizontalHeaderItem(i).text() for i in range(self.table.columnCount())]
        )
        for row in range(self.table.rowCount()):
            for col in range(self.table.columnCount()):
                item = self.table.item(row, col)
                preview_table.setItem(row, col, QTableWidgetItem(item.text() if item else ""))

        dialog = PrintPreviewDialog(preview_table, self)
        dialog.exec_()


##################################
#Constant
MainUI,_=loadUiType('interface2.ui')

department_choices = {
    1:"----",
    2:'المدير',
    3:'الشؤون الإدارية',
    4:'المعلوماتية',
    5:'الجباية',
    6:"الخزينة",
    7:"الدخل",
    8:"الاستعلام الضريبي",
    9:"الواردات",
    10:"كبار المكلفين"
}
Permision_List = [ ('a','اضافة موظف'),('b',"حذف موظف"),('c',"تعديل موظف"),('d',"اضافة دائرة"),('e',"حذف دائرة"),           
    ('f',"اضافة بريد"),('g',"حذف بريد"),('h',"تعديل بريد"),('i',"مراقبة العمل"),
    ('i',"إضافة إلى الاختيارات"),('k',"حذف من الاختيارات"),('l',"الحركة اليومية"),('m',"التقارير"),('n',"فلترة")]
Permision_Dict= {
    'a':'اضافة موظف'
    ,'b':"حذف موظف"
    ,'c':"تعديل موظف"
    ,'d':"اضافة دائرة"
    ,'e':"حذف دائرة"
    ,'f':"اضافة بريد"
    ,'g':"حذف بريد"
    ,'h':"تعديل بريد"
    ,'i':"إضافة إلى الاختيارات"
    ,'k':"حذف من الاختيارات"
    ,'l':"الحركة اليومية"
    ,'m':"التقارير"
    ,'n':"فلترة"
}
    
mail_type_tuble = [(1,'صادر خارجي'),(2,'صادر داخلي'),(3,'صادر هيئة'),(4,'صادر وزارة'),(5,'قرارات المديرية'),(6,'قرارات الطي'),
                   (7,'وارد خارجي'),(8,'وارد داخلي'),(9,'وارد وزارة'),(10,'وارد هيئة'),(11,"سجل الأحكام القضائية"),(12,'قيد مالي'),
                   (13,'الاعتراضات'),(14,'رواتب وأجور'),(15,'طلبات قسم الدخل'),(16,'طلبات عدا الدخل'),
                   (17,'مراسلات أقسام'),(18,'غير ذلك')]
mail_type_dict= {
                1:'صادر خارجي',
                2:'صادر داخلي',
                3:'صادر هيئة',
                4:'صادر وزارة',
                5:'قرارات المديرية',
                6:'قرارات الطي',
                7:'وارد خارجي',
                8:'وارد داخلي',
                9:'وارد وزارة',
                10:'وارد هيئة',
                11:"سجل الأحكام القضائية",
                12:'قيد مالي',
                13:'الاعتراضات',
                14:'رواتب وأجور',
                15:'طلبات قسم الدخل',
                16:'طلبات عدا الدخل',
                17:'مراسلات أقسام',
                18:'غير ذلك'}

process_name_list = [(1,'تسجيل الدخول'),(2,"تسجيل خروج"),(3,"تغيير كلمة السر"),(4,"ادخال بريد"),(5,"حذف بريد"),
                     (6,"تعديل بريد"),(7,"فلترة"),(8,"تقارير"),(9,"تعديل باسوورد ادمن"),(10,"إضافة دائرة"),
                     (11,"حذف دائرة"),(12,"إضافة موظف"),(13,"تعديل موظف"),(14,"حذف موظف"),(15,"تهيئة الاختيارات")]
process_name_dict = {
    1: 'تسجيل الدخول',
    2: 'تسجيل خروج',
    3: 'تغيير كلمة السر',
    4: 'ادخال بريد',
    5: 'حذف بريد',
    6: 'تعديل بريد',
    7: 'فلترة',
    8: 'تقارير',
    9: 'تعديل باسوورد ادمن',
    10: 'إضافة دائرة',
    11: 'حذف دائرة',
    12: 'إضافة موظف',
    13: 'تعديل موظف',
    14: 'حذف موظف',
    15: 'إضافة صادر من',
    16: 'إضافة صادر إلى',
    17: 'إضافة وارد من',
    18: 'إضافة وارد إلى',
    19: 'إضافة مستلم البريد ',
    20: 'إضافة الجهة المرسل إليها ',
    21: 'إضافة نوع الحكم ',
    22: 'حذف صادر من',
    23: 'حذف صادر إلى',
    24: 'حذف وارد من',
    25: 'حذف وارد إلى',
    26: 'حذف مستلم البريد ',
    27: 'حذف الجهة المرسل إليها ',
    28: 'حذف نوع الحكم '
    
}

def get_key_by_value(dictionary, value):
    for key, val in dictionary.items():
        if val == value:
            # print(key)
            return key
    return None  # إذا ما وجد القيمة
def get_value_by_key(dictionary, keyv):
    for key, val in dictionary.items():
        if keyv == key:
            # print(val)
            return val
    return None  # إذا ما وجد القيمة

employee_id = 0
employee_name = ''
employee_id_admin  = 0
employee_name_list = []

# Create QDate object
########################################
#####//class for ui
class Main(QMainWindow,MainUI):
    
    def __init__(self, parent=None):
        super(Main, self).__init__(parent)
        # QMainWindow.__init__(self)
        self.setupUi(self)
        self.setWindowFlags(QtCore.Qt.FramelessWindowHint)
        self.setWindowIcon(QtGui.QIcon(u":/icons/png/tax.png"))
        self.setWindowTitle("برنامج الأرشفة الالكترونية")
        # self.size_grip = QSizeGrip(self)
        QSizeGrip(self.size_grip)

        self.icon_show = QIcon("png/eye_open.png")
        self.icon_hide = QIcon("png/eye_closed.png")

        # إنشاء Action للزر
        self.toggle_action = QAction(self.icon_hide, "إظهار/إخفاء", self)
        self.toggle_action.triggered.connect(self.toggle_password)

        self.verticalSlider.setMinimum(600)
        self.verticalSlider.setMaximum(1200)
        self.verticalSlider.setValue(600)

        self.horizontalSlider.setMinimum(1008)
        self.horizontalSlider.setMaximum(1600)
        self.horizontalSlider.setValue(1008)

        self.dial.setMinimum(8)
        self.dial.setMaximum(18)
        self.dial.setValue(8)
        self.lcdNumber.display(8)
             
        self.label_2.mouseMoveEvent = self.moveWindow

        self.current_img_byte = None
        self.current_image_path = None # لتخزين المسار
        self.current_qr_byte = None
        self.current_qr_data = None
        ########## Server
        self.server_process = None
        self.used_numbers = set() 

        # إدراج الزر داخل QLineEdit
        self.lineEdit.addAction(self.toggle_action, QLineEdit.TrailingPosition)
        
        self.UI_Changes()
        self.Handle_Buttons()
        self.DB_Connect()
        self.LineEdit()
     
        #ِAksam
        self.Init_4_17_6()
        
        self.comboBox_4.currentTextChanged.connect(self.Handle_combobox)
        self.comboBox_17.currentTextChanged.connect(self.Handle_combobox)
        self.comboBox_6.currentTextChanged.connect(self.Handle_combobox)
        self.comboBox_6.currentTextChanged.connect(self.Show_All_Circle_inTable)

        self.comboBox_4.setCurrentIndex(0)
        self.comboBox_17.setCurrentIndex(0)
        self.Handle_combobox(self.comboBox_4.currentIndex())
        self.Handle_combobox(self.comboBox_17.currentIndex())

        self.comboBox_7.currentIndexChanged.connect(self.Handle_comboBox_7)
        self.comboBox_8.currentIndexChanged.connect(self.FilterMails)
        self.Handle_comboBox_7(self.comboBox_7.currentIndex())
        # self.comboBox_2.currentIndexChanged.connect(lambda:self.HandlecomboBox_2())
        self.comboBox_9.currentIndexChanged.connect(lambda:self.Handle_Daily_Work9())
        self.comboBox_18.currentIndexChanged.connect(lambda:self.Handle_Daily_Work18())
        self.comboBox_33.currentTextChanged.connect(lambda:self.generate_report_content())
        #self.date_process.currentIndexChanged.connect(lambda:self.Handle_Daily_Work())
        # self.HandlecomboBox_2()
        #self.Show_All_Circle_inTable(self.comboBox_6.currentText())
        
      
        self.Show_sader_from()
        self.Show_sader_to()
        self.Show_wared_from()
        self.Show_wared_to()
        self.Show_mostlem()
        self.Show_morsel()
        self.Show_ahkam()
        
        #الجداول
        self.tableWidget.cellClicked.connect(self.get_data)
        self.createpermision()
        self.Initilaize_Operation()
        self.Show_All_Employee()
        self.Show_All_Mails()

        
    
        self.lineEdit_20.textChanged.connect(self.check_mobile_live)
        self.lineEdit_5.textChanged.connect(self.check_email_live)

        ####################
#################################

##########################################
    ####//moving windows
    def moveWindow(self,e):
        if  self.isMaximized() == False:
            if e.buttons() == Qt.LeftButton:
                self.move(self.pos() + e.globalPos()-self.clickPosition)
                self.clickPosition = e.globalPos()
                e.accept()       

    def toggle_password(self):
        if self.lineEdit.echoMode() == QLineEdit.Password:
            self.lineEdit.setEchoMode(QLineEdit.Normal)
            self.toggle_action.setIcon(self.icon_show)
        else:
            self.lineEdit.setEchoMode(QLineEdit.Password)
            self.toggle_action.setIcon(QIcon("eye_closed.png"))

    def closewin(self):
        global employee_id
        if employee_id == 0:
            self.db.close()
            self.close()
        else:
            logoutid = 2
            nameprocess = get_value_by_key(process_name_dict,logoutid)
            
            # print(employee_id)
            datee = datetime.now()
            try:
                self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                self.db.commit()
            except Exception as ex:
                pass
            self.db.close()
            self.close()
            
    #progressBar
    def startProgress(self):
        self.step = 0
        self.progressBar.setValue(0)
        self.timer.start(100, self)

    def timerEvent(self, event):
        if self.step >= 100:
            self.timer.stop()
            # Optionally trigger login success logic here
            return
        self.step += 10
        self.progressBar.setValue(self.step)

    def mousePressEvent(self,event):
        self.clickPosition = event.globalPos()
            
    def restore_or_maximize_window(self):
        if self.isMaximized():
            self.showNormal()
            self.mini_max_window.setIcon(QtGui.QIcon(u":/icons/png/maximizew.png"))
        else:
            self.showMaximized()
            self.mini_max_window.setIcon(QtGui.QIcon(u":/icons/png/minimizew.png"))

    def Auto_Complete(self,model):
        model.setStringList(employee_name_list)
    
    def LineEdit(self):
        namelineedit = self.lineEdit_2
        completer = QCompleter()
        namelineedit.setCompleter(completer)
        model = QStringListModel()
        completer.setModel(model)
        self.Auto_Complete(model)

    def Init_4_17_6(self):
        self.cur.execute('''SELECT name FROM department ORDER By id ''')
        dep_name = self.cur.fetchall()
        # print(dep_name)
        for name in dep_name:
            self.comboBox_4.addItem(name[0])  
            self.comboBox_17.addItem(name[0]) 
            self.comboBox_6.addItem(name[0]) 

    def check_mobile_live(self):
        number = self.lineEdit_20.text().strip()

        pattern_international = r"^\+963\d{9}$"
        pattern_local = r"^09\d{8}$"

        if re.match(pattern_international, number) or re.match(pattern_local, number):
            self.lineEdit_20.setStyleSheet("border: 2px solid green;")
            self.statusBar().showMessage("✅ رقم الهاتف صالح")
        else:
            self.lineEdit_20.setStyleSheet("border: 2px solid red;")
            self.statusBar().showMessage("❌ رقم الهاتف غير صالح. يجب أن يبدأ بـ 09 أو +963 ويتكون من 10 أرقام")
    
    def check_email_live(self):
        email = self.lineEdit_5.text().strip()
        pattern = r'^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$'

        if re.match(pattern, email):
            self.lineEdit_5.setStyleSheet("border: 2px solid green;")
            self.statusBar().showMessage("✅ البريد الإلكتروني صالح")
        else:
            self.lineEdit_5.setStyleSheet("border: 2px solid red;")
            self.statusBar().showMessage("❌ البريد الإلكتروني غير صالح. تأكد من التنسيق مثل example@domain.com")

#######################################
####//Proccesures     
    def slideRightMenu(self):
        width = self.rightbody.width()
        if width == 55 :
            newWidth = 200
        else:
            newWidth = 55
        self.animation = QPropertyAnimation(self.rightbody,b"minimumWidth")
        self.animation.setDuration(250)
        self.animation.setStartValue(width)
        self.animation.setEndValue(newWidth)
        self.animation.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
        self.animation.start()
        
    def slideLeftMenu(self):
        width = self.leftbody.width()
        if width == 0 :
           # print(width)
            newWidth = 200
        else:
            newWidth = 0
        self.animation = QPropertyAnimation(self.leftbody,b"minimumWidth")
        self.animation.setDuration(250)
        self.animation.setStartValue(width)
        self.animation.setEndValue(newWidth)
        self.animation.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
        self.animation.start()

    def createpermision(self):

        per_str = ''
        if self.checkBox_5.isChecked():
            per_str = "".join([per_str, "abcdefghiklmn"])
        else :
            #1 اضافة موظف
            if self.checkBox_26.isChecked():
                per_str = "".join([per_str, "a"])
                #print(per_str)
            else:
                per_str.replace('a','')
                #print(per_str)
            
            #حذف موظف
            if self.checkBox.isChecked():
                per_str = "".join([per_str, "b"])
                #print(per_str)
            else:
                per_str.replace('b','')
                #print(per_str)
            
            #تحديث موظف
            if self.checkBox_2.isChecked():
                per_str = "".join([per_str, "c"])
                #print(per_str)
            else:
                per_str.replace('c','')
                #print(per_str)
            #اضافة دائرة
            if self.checkBox_29.isChecked():
                per_str = "".join([per_str, "d"])
                #print(per_str)
            else:
                per_str.replace('d','')
                #print(per_str)
            #حذف دائرة
            if self.checkBox_3.isChecked():
                per_str = "".join([per_str, "e"])
                #print(per_str)
            else:
                per_str.replace('e','')
                #print(per_str)
            # اضافة بريد 
            if self.checkBox_34.isChecked():
                per_str = "".join([per_str, "f"])
                #print(per_str)
            else:
                per_str.replace('f','')
                #print(per_str)
            #حذف بريد
            if self.checkBox_35.isChecked():
                per_str = "".join([per_str, "g"])
                #print(per_str)
            else:
                per_str.replace('g','')
            # print(per_str)
            #تعديل بريد 
            if self.checkBox_36.isChecked():
                per_str = "".join([per_str, "h"])
                #print(per_str)
            else:
                per_str.replace('h','')
                #print(per_str)

            #تهيئة الاختيارات 
            if self.checkBox_30.isChecked():
                per_str = "".join([per_str, "i"])
                #print(per_str)
            else:
                per_str.replace('i','')
                #print(per_str)

            #حذف  الاختيارات 
            if self.checkBox_37.isChecked():
                per_str = "".join([per_str, "k"])
                #print(per_str)
            else:
                per_str.replace('k','')
                #print(per_str)

            #حركة يومية
            if self.checkBox_31.isChecked():
                per_str = "".join([per_str, "l"])
                #print(per_str)
            else:
                per_str.replace('l','')
                #print(per_str)

            #تقارير
            if self.checkBox_32.isChecked():
                per_str = "".join([per_str, "m"])
                #print(per_str)
            else:
                per_str.replace('m','')
                #print(per_str)

            #فلترة
            if self.checkBox_7.isChecked():
                per_str = "".join([per_str, "n"])
                #print(per_str)
            else:
                per_str.replace('n','')
                #print(per_str)

        
        
        return per_str
        
    def UI_Changes(self):

        self.tabWidget.setEnabled(True)
        self.tabWidget_2.setEnabled(False)

        self.tabWidget.setCurrentWidget(self.login)
        self.tabWidget_2.setCurrentWidget(self.tab21)
        self.tabWidget.tabBar().setVisible(False)
        self.tabWidget_2.tabBar().setVisible(False)

        self.checkBox_26.stateChanged.connect(self.createpermision)
        self.checkBox.stateChanged.connect(self.createpermision)
        self.checkBox_2.stateChanged.connect(self.createpermision)
        self.checkBox_29.stateChanged.connect(self.createpermision)
        self.checkBox_3.stateChanged.connect(self.createpermision)
        self.checkBox_34.stateChanged.connect(self.createpermision)
        self.checkBox_35.stateChanged.connect(self.createpermision)
        self.checkBox_36.stateChanged.connect(self.createpermision)
        self.checkBox_30.stateChanged.connect(self.createpermision)
        self.checkBox_37.stateChanged.connect(self.createpermision)
        self.checkBox_31.stateChanged.connect(self.createpermision)
        self.checkBox_32.stateChanged.connect(self.createpermision)
        self.checkBox_7.stateChanged.connect(self.createpermision)
        
        self.dateEdit_12.setDate(QDate.currentDate())
        self.dateEdit_12.setDisplayFormat("yyyy-MM-dd")
        self.dateEdit_3.setDate(QDate.currentDate())
        self.dateEdit_3.setDisplayFormat("yyyy-MM-dd")
        self.dateEdit_2.setDate(QDate.currentDate())
        self.dateEdit_2.setDisplayFormat("yyyy-MM-dd")
        self.dateEdit_14.setDate(QDate.currentDate())
        self.dateEdit_14.setDisplayFormat("yyyy-MM-dd")
        self.dateEdit_13.setDate(QDate.currentDate())
        self.dateEdit_13.setDisplayFormat("yyyy-MM-dd")
        self.dateEdit_16.setDate(QDate.currentDate())
        self.dateEdit_16.setDisplayFormat("yyyy-MM-dd")
        self.dateEdit_15.setDate(QDate.currentDate())
        self.dateEdit_15.setDisplayFormat("yyyy-MM-dd")

        global employee_name,employee_id


        self.fontComboBox.currentFontChanged.connect(lambda _: self.Handle_Font_Change(employee_name))
        self.dial.valueChanged.connect(lambda _: self.Handle_Font_Change(employee_name))

        self.verticalSlider.valueChanged.connect(lambda _:self.update_window_size())
        self.horizontalSlider.valueChanged.connect(lambda _:self.update_window_size())

######################################## 
######//Database  
    def load_db_config(self):
        try:
            load_dotenv(dotenv_path="DB.env")  # تحميل القيم من ملف .env
        except Exception as s:
            self.statusBar().showMessage("ملف تكوين الاتصال  غير موجود")

        config = {
            "host": os.getenv("DB_HOST"),
            "user": os.getenv("DB_USER"),
            "password": os.getenv("DB_PASSWORD"),
            "database": os.getenv("DB_NAME")
        }
        return config

    def connect_to_db(self):
        config = self.load_db_config()
        try:
            self.db = mysql.connector.connect(**config)
            self.textEdit.setPlainText("نجحت عملية إختبار الاتصال يمكنك البدء في العمل  إبدأ بانشاء قاعدة البيانات إن لم تنشئها سابقا")
            self.statusBar().showMessage("تم تأسيس الاتصال")
            # statusBar().showMessage("تم تأسيس الاتصال")
            return self.db
        except Exception as e:
            self.textEdit.setPlainText(f"لم تنجح عملية الاتصال بقاعدة البيانات لوجود الخطأ التالي :{e}  يرجى ارسال النتيجة لقسم المعلوماتية")
            return None

    def DB_Connect(self):        
      
        self.db = self.connect_to_db()
        if self.db == None :
            self.statusBar().showMessage("لم يتم تأسيس الاتصال تأكد من اتصالك بالشبكة")
        else:
            self.statusBar().showMessage("تم تأسيس الاتصال")
            self.cur = self.db.cursor()
    
        self.cur.execute('''SELECT name FROM employees''')
        names = self.cur.fetchall()
        for name in names:
            employee_name_list.append(name[0])
        # print(employee_name_list)

    def create_backup(self):
        self.sql_file = ""
        self.zip_file = ""
        self.config = self.load_db_config()
      

        # استخراج القيم
        host = os.getenv("DB_HOST")
        user = os.getenv("DB_USER")
        password = os.getenv("DB_PASSWORD")
        db_name = os.getenv("DB_NAME")

        self.folder = QFileDialog.getExistingDirectory(self, "اختر مجلد الحفظ")
        if not self.folder:
            return

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.sql_file = os.path.join(self.folder, f"{db_name}_backup_{timestamp}.sql")

        dump_command = f"mysqldump -h {host} -u {user} -p{password} {db_name} > \"{self.sql_file}\""
        result = os.system(dump_command)

        if result == 0:
            QMessageBox.information(self, "تم", f"تم إنشاء النسخة:\n{self.sql_file}")
        else:
            QMessageBox.critical(self, "خطأ", "فشل في إنشاء النسخة الاحتياطية.")

    def compress_backup(self):
        if not self.sql_file or not os.path.exists(self.sql_file):
            QMessageBox.warning(self, "تنبيه", "لم يتم العثور على ملف النسخة الاحتياطية.")
            return

        self.zip_file = self.sql_file.replace(".sql", ".zip")
        try:
            with zipfile.ZipFile(self.zip_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
                zipf.write(self.sql_file, os.path.basename(self.sql_file))
            os.remove(self.sql_file)
            QMessageBox.information(self, "تم", f"تم ضغط الملف:\n{self.zip_file}")
        except Exception as e:
            QMessageBox.critical(self, "خطأ في الضغط", str(e))
    
    def restore_mysql(self):
        self.sql_file = ""
        self.zip_file = ""
        self.config = self.load_db_config()
      

        # استخراج القيم
        host = os.getenv("DB_HOST")
        user = os.getenv("DB_USER")
        password = os.getenv("DB_PASSWORD")
        db_name = os.getenv("DB_NAME")

        # اختيار ملف النسخة الاحتياطية
        sql_file, _ = QFileDialog.getOpenFileName(self, "اختر ملف النسخة الاحتياطية", "", "SQL Files (*.sql)")
        if not sql_file:
            return
        
        # تأكيد من المستخدم قبل الاسترجاع
        reply = QMessageBox.question(
            self,
            "تأكيد الاسترجاع",
            f"⚠️ هل أنت متأكد أنك تريد استرجاع النسخة الاحتياطية؟\nسيتم الكتابة فوق قاعدة البيانات الحالية: {db_name}",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )

        if reply != QMessageBox.Yes:
            return  # المستخدم ألغى العملية

        # تنفيذ أمر الاسترجاع
        restore_command = f"mysql -h {host} -u {user} -p{password} {db_name} < \"{sql_file}\""
        result = os.system(restore_command)

        if result == 0:
            QMessageBox.information(self, "تم", f"✅ تم استرجاع النسخة بنجاح من:\n{sql_file}")
        else:
            QMessageBox.critical(self, "خطأ", "❌ فشل في استرجاع النسخة الاحتياطية.")
     
##################################

    def toggle_server(self):
        if self.server_process is None:
            self.start_server()
        else:
            self.stop_server()
    
    def start_server(self):
        self.server_process = QProcess()
        self.server_process.start('python', ['server.py'])
        self.server_btn.setText("إيقاف الخادم")
        #QMessageBox.information(self, "تم التشغيل", "الخادم جاهز لاستقبال الملفات على المنفذ 5000")
        self.statusBar().showMessage("الخادم جاهز لاستقبال الملفات")
    
    def stop_server(self):
        if self.server_process:
            self.server_process.terminate()
            self.server_process.waitForFinished(3000)  # انتظر حتى تنتهي العملية خلال 3 ثوانٍ
            self.server_process = None
            self.server_btn.setText("تشغيل خادم الاستقبال")
            self.statusBar().showMessage("❌ تم إيقاف خادم الاستقبال")
        else:
            self.statusBar().showMessage("⚠️ لا يوجد خادم يعمل حالياً")

    def show_uploaded_files(self):
        try:
            with open('upload_log.json', 'r', encoding='utf-8') as f:
                logs = json.load(f)

            self.textEdit.clear()
            for entry in logs:
                line = (
                    f"📁 الملف: {entry['filename']}\n"
                    f"🌐 IP المرسل: {entry['ip']}\n"
                    f"📦 الحجم: {entry['size_MB']} MB\n"
                    f"🕒 التاريخ: {entry['timestamp']}\n"
                    f"📍 المسار: {entry['path']}\n"
                    "-----------------------------\n"
                )
                self.textEdit.append(line)

        except Exception as e:
            self.textEdit.setPlainText(f"❌ خطأ في قراءة الملفات:\n{str(e)}")

#####################################################  
##########REPORTS###############
    
    def preview_and_print_report(self):
        printer = QPrinter(QPrinter.HighResolution)
        preview_dialog = QPrintPreviewDialog(printer, self)
        preview_dialog.setWindowTitle("معاينة قبل الطباعة")
        preview_dialog.paintRequested.connect(self.preview.print_)
        preview_dialog.exec_()

    def export_report_to_word(self):
        nameemp = self.comboBox_33.currentText()

        self.cur.execute("""
            SELECT dailymovements.process_type, dailymovements.date_process, employees.name
            FROM dailymovements
            JOIN employees ON dailymovements.emp_id = employees.id
            WHERE employees.name = %s
        """, [nameemp])
        
        result = self.cur.fetchall()

        if result:
            empname = result[0][2]
            archived = len(result)

            # فتح مربع حفظ الملف
            file_path, _ = QFileDialog.getSaveFileName(self, "حفظ التقرير", "", "Word Files (*.docx)")
            if not file_path:
                return

            # إنشاء المستند
            doc = Document()
            doc.add_picture("3kab.png", width=Inches(1.0))
            doc.add_heading("الجمهورية العربية السورية", level=1)
            doc.add_heading(f"تقرير الموظف: {empname}", level=2)
            doc.add_paragraph("هذا التقرير يعرض ملخصًا للعمليات التي قام بها الموظف المختار.")
            doc.add_paragraph(f"عدد العمليات المؤرشفة: {archived}")
            doc.add_heading("تفاصيل العمليات:", level=3)

            # جدول العمليات
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'نوع العملية'
            hdr_cells[1].text = 'تاريخ العملية'

            for operation_type, date_operation, _ in result:
                label = process_name_dict.get(operation_type, "عملية غير معروفة")
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = str(date_operation)

            doc.save(file_path)

    def generate_report_content(self):
        nameemp = self.comboBox_33.currentText()
        
        self.cur.execute("""SELECT dailymovements.process_type, dailymovements.date_process, employees.name
                         FROM dailymovements JOIN employees ON dailymovements.emp_id = employees.id
                         WHERE employees.name = %s""", [nameemp])
        result = self.cur.fetchall()
        #print(len(result))
        #print (result)
        if result:
            empname = result[0][2]  # اسم الموظف من أول سجل
            archived = len(result)

            # بناء جدول العمليات
            operations_html = """
            <table border='1' cellspacing='0' cellpadding='4' style='width: 100%; text-align: center;'>
                <tr><th>نوع العملية</th><th>تاريخ العملية</th></tr>
            """
            for operation_type, date_operation, _ in result:
                # print(f"نوع العملية المسترجعة: {operation_type} - الترجمة: {process_name_dict.get(operation_type, 'عملية غير معروفة')}")

                # operations_html += f"<tr><td>{process_name_dict.get(operation_type, "عملية غير معروفة")}</td><td>{date_operation}</td></tr>"
                operations_html += f"<tr><td>{operation_type}</td><td>{date_operation}</td></tr>"

            operations_html += "</table>"

            # بناء التقرير الكامل
            html = f"""
            <div style='text-align: center;'>
                <img src='3kab.png' width='100' height='100' />
                <h2>الجمهورية العربية السورية</h2>
                <h3>تقرير الموظف: {empname}</h3>
            </div>
            <p>هذا التقرير يعرض ملخصًا للعمليات التي قام بها الموظف المختار.</p>
            <ul>
                <li>عدد العمليات المؤرشفة: {archived}</li>
            </ul>
            <h4>تفاصيل العمليات:</h4>
            {operations_html}
            """

            self.preview.setHtml(html)
        else:
            self.preview.setHtml("<p>لا توجد عمليات لهذا الموظف.</p>")

#########################################################
#####//connect menu buttons with tab    

    def Handle_Buttons(self):
        #tab
        self.btndaily.clicked.connect(lambda:self.tabWidget.setCurrentWidget(self.dailymovement))
        self.btndaily.clicked.connect(lambda:self.Show_All_Operation(employee_id))
        self.btnsearch.clicked.connect(lambda:self.tabWidget.setCurrentWidget(self.search))
        self.btndelete.clicked.connect(lambda:self.tabWidget.setCurrentWidget(self.delete))
        self.btnreport.clicked.connect(lambda:self.tabWidget.setCurrentWidget(self.reports))
        self.btnsettings.clicked.connect(lambda:self.tabWidget.setCurrentWidget(self.setting))
        self.btnsettings.clicked.connect(lambda:self.tabWidget_2.setCurrentWidget(self.tab24))
        self.btnentry.clicked.connect(lambda:self.tabWidget.setCurrentWidget(self.entermail))
        self.pushButton_23.clicked.connect(lambda:self.tabWidget_2.setCurrentWidget(self.tab))
        self.pushButton_16.clicked.connect(lambda:self.tabWidget.setCurrentWidget(self.changepassword))
        self.pushButton_4.clicked.connect(lambda:self.tabWidget_2.setCurrentWidget(self.tab21))
        self.pushButton.clicked.connect(lambda:self.tabWidget_2.setCurrentWidget(self.tab22))
        self.pushButton_3.clicked.connect(lambda:self.tabWidget_2.setCurrentWidget(self.tab23))
        self.pushButton_2.clicked.connect(lambda:self.tabWidget_2.setCurrentWidget(self.tab24))
        self.pushButton_33.clicked.connect(lambda:self.tabWidget.setCurrentWidget(self.tab_2))
        self.pushButton_36.clicked.connect(lambda:self.tabWidget.setCurrentWidget(self.login))
        self.pushButton_10.clicked.connect(lambda:self.tabWidget_2.setCurrentWidget(self.tab_3))
        self.pushButton_56.clicked.connect(lambda:self.tabWidget_2.setCurrentWidget(self.tab_4))

        #procedure
        self.pushButton_15.clicked.connect(lambda:self.Handle_Login())
        global employee_name
        self.pushButton_50.clicked.connect(lambda:self.Load_User_Settings(employee_name))
        self.pushButton_49.clicked.connect(lambda: self.Save_All_User_Settings(employee_name))
        self.btnclose.clicked.connect(lambda:self.closewin())
        self.mini_window.clicked.connect(lambda:self.showMinimized())
        self.close_window.clicked.connect(lambda:self.closewin())
        self.mini_max_window.clicked.connect(lambda:self.restore_or_maximize_window())
        self.pushButton_28.clicked.connect(lambda:self.Add_New_Circle())
        self.pushButton_26.clicked.connect(lambda:self.Del_Circle())
        try:
            self.pushButton_21.clicked.connect(lambda:self.Add_New_Employee())
        except Exception as ex:
            self.statusBar().showMessage(f'{ex}')
        self.stsrtwindow.clicked.connect(lambda:self.slideRightMenu())
        self.helpwindow.clicked.connect(lambda:self.slideLeftMenu())
        self.pushButton_22.clicked.connect(lambda:self.clear_Cells())
        self.pushButton_14.clicked.connect(lambda:self.Search_Employee())
        self.pushButton_20.clicked.connect(lambda:self.Update_Employee())
        self.pushButton_24.clicked.connect(lambda:self.Del_Employee())
        self.pushButton_35.clicked.connect(lambda:self.Add_New_Mail())
        self.pushButton_25.clicked.connect(lambda:self.Del_Mail())
        self.pushButton_18.clicked.connect(lambda:self.Search_Mail())
        self.pushButton_34.clicked.connect(lambda:self.Clear_Data_Add())
        self.pushButton_11.clicked.connect(lambda:self.Add_Text_To_sader_from())
        self.pushButton_12.clicked.connect(lambda:self.Add_Text_To_sader_to())
        self.pushButton_27.clicked.connect(lambda:self.Add_Text_To_wared_from())
        self.pushButton_29.clicked.connect(lambda:self.Add_Text_To_wared_to())
        self.pushButton_30.clicked.connect(lambda:self.Add_Text_To_mostlem())
        self.pushButton_31.clicked.connect(lambda:self.Add_Text_To_morsel())
        self.pushButton_32.clicked.connect(lambda:self.Add_Text_To_cb_ahkam())
        self.pushButton_9.clicked.connect(lambda:self.Update_mail())
        self.pushButton_17.clicked.connect(lambda:self.ChangeUserPass())
        self.pushButton_19.clicked.connect(lambda:self.ChangeAdminPass())
        #filter
        self.pushButton_37.clicked.connect(lambda:self.Export_xlsx2())
        self.pushButton_38.clicked.connect(lambda:self.Export_cvx())
        #dailymovment
        self.pushButton_39.clicked.connect(lambda:self.Export_Operation_xls())
        self.pushButton_40.clicked.connect(lambda:self.Export_Operation_cvs())
        self.pushButton_41.clicked.connect(lambda:self.on_print_button_clicked())

        #report
        self.btnPrint.clicked.connect(self.preview_and_print_report)
        self.btnExport.clicked.connect(self.export_report_to_word)
        
        self.pushButton_42.clicked.connect(self.del_sader_from)
        self.pushButton_45.clicked.connect(self.del_sader_to)
        self.pushButton_44.clicked.connect(self.del_wared_from)
        self.pushButton_43.clicked.connect(self.del_wared_to)
        self.pushButton_46.clicked.connect(self.del_mostlem)
        self.pushButton_47.clicked.connect(self.del_morsel)
        self.pushButton_48.clicked.connect(self.del_ahkam)
        ###########database
        self.pushButton_53.clicked.connect(self.connect_to_db)
        self.pushButton_52.clicked.connect(self.create_backup)
        self.pushButton_54.clicked.connect(self.compress_backup)
        # self.pushButton_56.clicked.connect(self.send_backup_file)
        self.pushButton_55.clicked.connect(self.restore_mysql)
        self.pushButton_58.clicked.connect(self.show_uploaded_files)
        self.server_btn.clicked.connect(self.toggle_server)
        self.server_btn.clicked.connect(self.stop_server)
        

        #progressBar
        self.pushButton_15.clicked.connect(lambda:self.startProgress())
        self.timer = QBasicTimer()
        self.step = 0
        ########Images
        self.pushButton_6.clicked.connect(lambda:self.BrowseImage())
        self.pushButton_5.clicked.connect(self.open_capture_window)
        self.pushButton_57.clicked.connect(lambda :self.open_qr_window())
        self.pushButton_59.clicked.connect(lambda :self.Search_QR())
        self.pushButton_60.clicked.connect(lambda :self.Print_QR())
    
    def Handle_Font_Change(self, username):
        font_name = self.fontComboBox.currentFont().family()
        font_size = self.dial.value()

        # إنشاء إعدادات باسم المستخدم
        settings = QSettings("HassanApps", "FontSizeApp")
        settings.setValue(f"{username}_font_family", font_name)
        settings.setValue(f"{username}_font_size", font_size)

        # تطبيق الخط
        font = QFont(font_name, font_size)
        for widget in self.findChildren(QWidget):
            widget.setFont(font)

        # عرض الحجم
        self.lcdNumber.display(font_size)

        print(f"تم تطبيق وحفظ خط المستخدم {username}: {font_name} بحجم {font_size}")

    def Load_User_Font_Settings(self, username):
        settings = QSettings("HassanApps", "FontSizeApp")
        font_name = settings.value(f"{username}_font_family", "Arial")
        font_size = settings.value(f"{username}_font_size", 8, type=int)

        # ضبط المكونات
        self.fontComboBox.setCurrentFont(QFont(font_name))
        self.dial.setValue(font_size)
        self.lcdNumber.display(font_size)

        # تطبيق الخط
        font = QFont(font_name, font_size)
        for widget in self.findChildren(QWidget):
            widget.setFont(font)

        print(f"تم استرجاع إعدادات الخط لـ {username}: {font_name} بحجم {font_size}")

    def apply_font_to_all_widgets(self, size, font_name):
        font = QFont(font_name, size)
        for widget in self.findChildren(QWidget):
            widget.setFont(font)

    def Save_All_User_Settings(self, username):
        font_name = self.fontComboBox.currentFont().family()
        font_size = self.dial.value()
        # theme_name = "dark_teal.xml"  # أو الثيم النشط حاليًا

        # الحصول على الحجم والموقع
        window_width = self.width()
        window_height = self.height()
        window_x = self.x()
        window_y = self.y()

        settings_file = "user_preferences.json"
        new_settings = {
            "username": username,
            "font_family": font_name,
            "font_size": font_size,
            "window_width": window_width,
            "window_height": window_height,
            "window_x": window_x,
            "window_y": window_y
        }

        all_settings = {}
        if os.path.exists(settings_file):
            with open(settings_file, "r", encoding="utf-8") as f:
                all_settings = json.load(f)

        # تحديث أو إضافة الإعدادات
        all_settings[username] = new_settings

        with open(settings_file, "w", encoding="utf-8") as f:
            json.dump(all_settings, f, indent=4, ensure_ascii=False)

        self.lcdNumber.display(font_size)
        self.apply_font_to_all_widgets(font_size, font_name)

        self.statusBar().showMessage(f"✅ تم حفظ الإعدادات الخاصة بـ {username} في ملف خارجي")
   
    def Load_User_Settings(self, username):
        settings_file = "user_preferences.json"
        if not os.path.exists(settings_file):
            return

        with open(settings_file, "r", encoding="utf-8") as f:
            all_settings = json.load(f)

        user_settings = all_settings.get(username)
        if not user_settings:
            return

        # تطبيق الحجم والموقع
        self.resize(user_settings["window_width"], user_settings["window_height"])
        self.move(user_settings["window_x"], user_settings["window_y"])

        # تطبيق الخط والثيم
        font = QFont(user_settings["font_family"], user_settings["font_size"])
        self.setFont(font)
        self.lcdNumber.display(user_settings["font_size"])

    def update_window_size(self):
        height = self.verticalSlider.value()
        width = self.horizontalSlider.value()
        self.resize(width, height)

###########################################
##Handle ComboBox      

    def getIntCircleID(self,d,c):
        self.cur.execute("""SELECT id FROM circle WHERE name=%s AND circle_department_id=%s""",[c,d])
        nameCirid = self.cur.fetchone()
        return nameCirid[0]

    def Handle_combobox(self,selected_item):
        #print(selected_item)
        self.comboBox_5.clear()
        self.comboBox_32.clear()
        self.comboBox_5.insertItem(0,"----")
        self.comboBox_32.insertItem(0,'----')
        if selected_item != 0 :
            #print(selected_item)
            # idc = Depart_List.index(selected_item)
            idc =get_key_by_value(department_choices,selected_item)
            #print(idc)


            if idc != 1 :
                self.cur.execute("""
                            SELECT name FROM circle WHERE circle_department_id = %s
                            """,[idc])
                circles = self.cur.fetchall()
                print(circles)
                for circle in circles:
                    
                    
                    self.comboBox_5.addItem(circle[0])
                    
                    
                    self.comboBox_32.addItem(circle[0])
            else:
                pass

####################
#For Type Mail
    
    def Handle_comboBox_7(self,selected_item):
        
        if selected_item != 0 :
            self.label_23.setText(f"{self.comboBox_7.currentText()}")
            self.Box1.setEnabled(True)
            self.Box4.setEnabled(True)
            self.Box9.setEnabled(True)
            self.Box13.setEnabled(True)
            self.Box15.setEnabled(False)
            self.Box5.setEnabled(True)

            ###sader    
            if selected_item == 1:
                self.Box2.setEnabled(True)
                self.Box8.setEnabled(False)
                self.Box3.setEnabled(False)

                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box14.setEnabled(False)

            elif selected_item == 2:
                self.Box2.setEnabled(True)
                self.Box8.setEnabled(False)
                self.Box3.setEnabled(False)
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box14.setEnabled(False)

            elif selected_item == 3:
                self.Box2.setEnabled(True)
                self.Box8.setEnabled(False)
                self.Box3.setEnabled(False)
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box14.setEnabled(False)

            elif selected_item == 4:
                self.Box2.setEnabled(True)
                self.Box8.setEnabled(False)
                self.Box3.setEnabled(False)
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box14.setEnabled(False)

            elif selected_item == 5:
                self.Box2.setEnabled(True)
                self.Box8.setEnabled(False)
                self.Box3.setEnabled(False)
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box14.setEnabled(False)

            elif selected_item == 6:
                self.Box2.setEnabled(True)
                self.Box3.setEnabled(False)
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box14.setEnabled(False)
            ###wared    
            elif selected_item == 7:
                self.Box3.setEnabled(True)
                self.Box2.setEnabled(False)
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box14.setEnabled(False)

            elif selected_item == 8:
                self.Box3.setEnabled(True)
                self.Box2.setEnabled(False)
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box14.setEnabled(False)

            elif selected_item == 9:
                self.Box3.setEnabled(True)
                self.Box2.setEnabled(False)
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box14.setEnabled(False)

            elif selected_item == 10:
                self.Box3.setEnabled(True)
                self.Box2.setEnabled(False)
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box14.setEnabled(False)

            ###ahkam
            elif selected_item == 11:
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)

                self.Box14.setEnabled(True)
                self.Box2.setEnabled(False)
                
                self.Box3.setEnabled(True)
                self.Box8.setEnabled(False)
                self.Box15.setEnabled(True)
                
            ##keed mali
            elif selected_item == 12:
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box2.setEnabled(True)
                self.Box3.setEnabled(True)
                self.Box14.setEnabled(False)
                self.Box8.setEnabled(False)
            ##mo3tred
            elif selected_item == 13:
                self.Box8.setEnabled(True)
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box3.setEnabled(True)
                self.Box2.setEnabled(False)
                self.Box14.setEnabled(False)

            ##rawateb
            elif selected_item == 14:
                self.Box8.setEnabled(False)
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box3.setEnabled(True)
                self.Box2.setEnabled(False)
                self.Box14.setEnabled(False)
                

            ##dakhel
            elif selected_item == 15:
                self.Box5.setEnabled(True)
                self.Box6.setEnabled(True)
                self.Box7.setEnabled(True)
                self.Box3.setEnabled(True)
                self.Box2.setEnabled(False)
                self.Box14.setEnabled(False)

            ##talabat
            elif selected_item == 16:
                self.Box5.setEnabled(True)
                self.Box3.setEnabled(True)
                self.Box2.setEnabled(False)
                self.Box14.setEnabled(False)

            ##moraslat
            elif selected_item == 17:
                self.Box2.setEnabled(True)
                self.Box3.setEnabled(True)
                self.Box14.setEnabled(False)
                self.Box5.setEnabled(False)
                self.Box6.setEnabled(False)
                self.Box7.setEnabled(False)

            ##other
            elif selected_item == 18:
                pass
        else :

            self.Box1.setEnabled(False)  
            self.Box2.setEnabled(False)  
            self.Box3.setEnabled(False)  
            self.Box4.setEnabled(False)  
            self.Box5.setEnabled(False)  
            self.Box6.setEnabled(False)  
            self.Box7.setEnabled(False)  
            self.Box8.setEnabled(False)  
            self.Box9.setEnabled(False)  
            self.Box13.setEnabled(False)  
            self.Box14.setEnabled(False)  


#######################        
    # INITILIZING COMBOBOX
    
    def Add_Text_To_sader_from(self): 
        name = self.lineEdit_51.text()
        if name!='':
            self.cur.execute("""INSERT INTO cb_sader_from (name) 
                    VALUES (%s)""",(name,))
        self.db.commit()
        self.statusBar().showMessage("تم إضافة البند بنجاح")
        self.lineEdit_51.setText('')
        self.Show_sader_from()
        addsaderfrom = 15
        nameprocess = get_value_by_key(process_name_dict,addsaderfrom)
        global employee_id
        #print(employee_id)
        datee = datetime.now()
        self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                        VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
        
        self.db.commit()

    def del_sader_from(self):
        sf = self.comboBox_26.currentText()
        if sf != "----" :
            reply = QMessageBox.critical(self, "تحذير", "هل تريد بالتأكيد الحذف?", QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.cur.execute('''
                             DELETE FROM cb_sader_from WHERE name=%s 
                             ''',[sf]) 
                self.db.commit()
                self.statusBar().showMessage("تم حذف البند بنجاح")
                self.Show_sader_from()
                delsaderfrom = 22
                nameprocess = get_value_by_key(process_name_dict,delsaderfrom)
                global employee_id
                #print(employee_id)
                datee = datetime.now()
                self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                
                self.db.commit()
            else:
                return

    def Show_sader_from(self):
        '''comboBox_14
        comboBox_22
        comboBox_26'''
        self.comboBox_14.clear()
        self.comboBox_22.clear()
        self.comboBox_26.clear()
        self.comboBox_14.addItem('----')
        self.comboBox_22.addItem('----')
        self.comboBox_26.addItem('----')

        self.cur.execute('''SELECT * FROM cb_sader_from''')
        data = self.cur.fetchall()
        #print(data)
        sorted_data = sorted(data, key=lambda x: x[0])
        for row in sorted_data:
            #print(row[0])
            
            self.comboBox_14.insertItem(self.comboBox_14.count(), row[1])
            self.comboBox_22.insertItem(self.comboBox_22.count(), row[1])
            self.comboBox_26.insertItem(self.comboBox_26.count(), row[1])
            
    def Add_Text_To_sader_to(self):
        name = self.lineEdit_52.text()
        if name!='':
            self.cur.execute("""INSERT INTO cb_sader_to (name) 
                    VALUES (%s)""",(name,))
        self.db.commit()
        self.statusBar().showMessage("تم إضافة البند بنجاح")
        self.lineEdit_52.setText('')
        self.Show_sader_to()
        addsaderfrom = 16
        nameprocess = get_value_by_key(process_name_dict,addsaderfrom)
        global employee_id
        #print(employee_id)
        datee = datetime.now()
        self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                        VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
        
        self.db.commit()

    def del_sader_to(self):
        sf = self.comboBox_27.currentText()
        if sf != "----" :
            reply = QMessageBox.critical(self, "تحذير", "هل تريد بالتأكيد الحذف?", QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.cur.execute('''
                             DELETE FROM cb_sader_to WHERE name=%s 
                             ''',[sf]) 
                self.db.commit()
                self.statusBar().showMessage("تم حذف البند بنجاح")
                self.Show_sader_to()
                delsaderfrom = 23
                nameprocess = get_value_by_key(process_name_dict,delsaderfrom)
                global employee_id
                #print(employee_id)
                datee = datetime.now()
                self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                
                self.db.commit()
            else:
                return

    def Show_sader_to(self):
        '''comboBox_15
        comboBox_19
        comboBox_27'''
        self.comboBox_15.clear()
        self.comboBox_19.clear()
        self.comboBox_27.clear()
        self.comboBox_15.addItem('----')
        self.comboBox_19.addItem('----')
        self.comboBox_27.addItem('----')
        self.cur.execute('''SELECT * FROM cb_sader_to''')
        data = self.cur.fetchall()
        #print(data)
        sorted_data = sorted(data, key=lambda x: x[0])
        for row in sorted_data:
            #print(row[0])
            self.comboBox_15.insertItem(self.comboBox_15.count(), row[1])
            self.comboBox_19.insertItem(self.comboBox_19.count(), row[1])
            self.comboBox_27.insertItem(self.comboBox_27.count(), row[1])            
            
    def Add_Text_To_wared_from(self):
        name = self.lineEdit_53.text()
        if name!='':
            self.cur.execute("""INSERT INTO cb_wared_from (name) 
                    VALUES (%s)""",(name,))
        self.db.commit()
        self.statusBar().showMessage("تم إضافة البند بنجاح")
        self.lineEdit_53.setText('')
        self.Show_wared_from()
        addwaredfrom = 17
        nameprocess = get_value_by_key(process_name_dict,addwaredfrom)
        global employee_id
        #print(employee_id)
        datee = datetime.now()
        self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                        VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
        
        self.db.commit()
        
    def Show_wared_from(self):
        '''comboBox_28
        comboBox_25
        comboBox_11'''
        self.comboBox_28.clear()
        self.comboBox_25.clear()
        self.comboBox_11.clear()
        self.comboBox_28.addItem('----')
        self.comboBox_25.addItem('----')
        self.comboBox_11.addItem('----')
        self.cur.execute('''SELECT * FROM cb_wared_from''')
        data = self.cur.fetchall()
        #print(data)
        sorted_data = sorted(data, key=lambda x: x[0])
        for row in sorted_data:
            #print(row[0])
            self.comboBox_28.insertItem(self.comboBox_28.count(), row[1])
            self.comboBox_25.insertItem(self.comboBox_25.count(), row[1])
            self.comboBox_11.insertItem(self.comboBox_11.count(), row[1])

    def del_wared_from(self):
        sf = self.comboBox_28.currentText()
        if sf != "----" :
            reply = QMessageBox.critical(self, "تحذير", "هل تريد بالتأكيد الحذف?", QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.cur.execute('''
                             DELETE FROM cb_wared_from WHERE name=%s 
                             ''',[sf]) 
                self.db.commit()
                self.statusBar().showMessage("تم حذف البند بنجاح")
                self.Show_wared_from()
                delsaderfrom = 24
                nameprocess = get_value_by_key(process_name_dict,delsaderfrom)
                global employee_id
                #print(employee_id)
                datee = datetime.now()
                self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                
                self.db.commit()
            else:
                return

    def Add_Text_To_wared_to(self):
        name = self.lineEdit_54.text()
        if name!='':
            self.cur.execute("""INSERT INTO cb_wared_to (name) 
                    VALUES (%s)""",(name,))
        self.db.commit()
        self.statusBar().showMessage("تم إضافة البند بنجاح")
        self.lineEdit_54.setText('')
        self.Show_wared_to()
        addsaderfrom = 18
        nameprocess = get_value_by_key(process_name_dict,addsaderfrom)
        global employee_id
        #print(employee_id)
        datee = datetime.now()
        self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                        VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
        
        self.db.commit()
        
    def Show_wared_to(self):
        '''comboBox_29
        comboBox_24
        comboBox_12'''
        self.comboBox_29.clear()
        self.comboBox_24.clear()
        self.comboBox_12.clear()
        self.comboBox_29.addItem('----')
        self.comboBox_24.addItem('----')
        self.comboBox_12.addItem('----')
        self.cur.execute('''SELECT * FROM cb_wared_to''')
        data = self.cur.fetchall()
        #print(data)
        sorted_data = sorted(data, key=lambda x: x[0])
        for row in sorted_data:
            #print(row[0])
            self.comboBox_29.insertItem(self.comboBox_29.count(), row[1])
            self.comboBox_24.insertItem(self.comboBox_24.count(), row[1])
            self.comboBox_12.insertItem(self.comboBox_12.count(), row[1])

    def del_wared_to(self):
        sf = self.comboBox_29.currentText()
        if sf != "----" :
            reply = QMessageBox.critical(self, "تحذير", "هل تريد بالتأكيد الحذف?", QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.cur.execute('''
                             DELETE FROM cb_wared_to WHERE name=%s 
                             ''',[sf]) 
                self.db.commit()
                self.statusBar().showMessage("تم حذف البند بنجاح")
                self.Show_wared_to()
                delsaderfrom = 25
                nameprocess = get_value_by_key(process_name_dict,delsaderfrom)
                global employee_id
                #print(employee_id)
                datee = datetime.now()
                self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                
                self.db.commit()
            else:
                return

    def Add_Text_To_mostlem(self):
        name = self.lineEdit_55.text()
        if name!='':
            self.cur.execute("""INSERT INTO cb_mostlem_bareed (name) 
                    VALUES (%s)""",(name,))
        self.db.commit()
        self.statusBar().showMessage("تم إضافة البند بنجاح")
        self.lineEdit_55.setText('')
        self.Show_mostlem()
        addsaderfrom = 19
        nameprocess = get_value_by_key(process_name_dict,addsaderfrom)
        global employee_id
        #print(employee_id)
        datee = datetime.now()
        self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                        VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
        
        self.db.commit()
        
    def Show_mostlem(self):
        '''comboBox_30
        comboBox_23
        comboBox_10'''
        self.comboBox_30.clear()
        self.comboBox_23.clear()
        self.comboBox_10.clear()
        self.comboBox_30.addItem('----')
        self.comboBox_23.addItem('----')
        self.comboBox_10.addItem('----')
        self.cur.execute('''SELECT name FROM cb_mostlem_bareed''')
        data = self.cur.fetchall()
        #print(data)
        #sorted_data = sorted(data, key=lambda x: x[0])
        for row in data:
            #print(row[0])
            self.comboBox_30.insertItem(self.comboBox_30.count(), row[0])
            self.comboBox_23.insertItem(self.comboBox_23.count(), row[0])
            self.comboBox_10.insertItem(self.comboBox_10.count(), row[0])

    def del_mostlem(self):
        sf = self.comboBox_30.currentText()
        if sf != "----" :
            reply = QMessageBox.critical(self, "تحذير", "هل تريد بالتأكيد الحذف?", QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.cur.execute('''
                             DELETE FROM cb_mostlem_bareed WHERE name=%s 
                             ''',[sf]) 
                self.db.commit()
                self.statusBar().showMessage("تم حذف البند بنجاح")
                self.Show_mostlem()
                delsaderfrom = 26
                nameprocess = get_value_by_key(process_name_dict,delsaderfrom)
                global employee_id
                #print(employee_id)
                datee = datetime.now()
                self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                
                self.db.commit()
            else:
                return

    def Add_Text_To_morsel(self):
        name = self.lineEdit_56.text()
        if name!='':
            self.cur.execute("""INSERT INTO cb_morsal_to (name) 
                    VALUES (%s)""",(name,))
        self.db.commit()
        self.statusBar().showMessage("تم إضافة البند بنجاح")
        self.lineEdit_56.setText('')
        self.Show_morsel()
        addsaderfrom = 20
        nameprocess = get_value_by_key(process_name_dict,addsaderfrom)
        global employee_id
        #print(employee_id)
        datee = datetime.now()
        self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                        VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
        
        self.db.commit()
        
    def Show_morsel(self):
        '''comboBox_31
        comboBox_20
        comboBox_13'''
        self.comboBox_31.clear()
        self.comboBox_20.clear()
        self.comboBox_13.clear()
        self.comboBox_31.addItem('----')
        self.comboBox_20.addItem('----')
        self.comboBox_13.addItem('----')
        self.cur.execute('''SELECT name FROM cb_morsal_to''')
        data = self.cur.fetchall()
        #print(data)
        #sorted_data = sorted(data, key=lambda x: x[0])
        for row in data:
            #print(row[0])
            self.comboBox_31.insertItem(self.comboBox_31.count(), row[0])
            self.comboBox_20.insertItem(self.comboBox_20.count(), row[0])
            self.comboBox_13.insertItem(self.comboBox_13.count(), row[0])

    def del_morsel(self):
        sf = self.comboBox_31.currentText()
        if sf != "----" :
            reply = QMessageBox.critical(self, "تحذير", "هل تريد بالتأكيد الحذف?", QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.cur.execute('''
                             DELETE FROM cb_morsal_to WHERE name=%s 
                             ''',[sf]) 
                self.db.commit()
                self.statusBar().showMessage("تم حذف البند بنجاح")
                self.Show_morsel()
                delsaderfrom = 27
                nameprocess = get_value_by_key(process_name_dict,delsaderfrom)
                global employee_id
                #print(employee_id)
                datee = datetime.now()
                self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                
                self.db.commit()
            else:
                return

    def Add_Text_To_cb_ahkam(self):
        name = self.lineEdit_57.text()
        if name!='':
            self.cur.execute("""INSERT INTO cb_ahkam (name) 
                    VALUES (%s)""",(name,))
        self.db.commit()
        self.statusBar().showMessage("تم إضافة البند بنجاح")
        self.lineEdit_57.setText('')
        self.Show_ahkam()
        addsaderfrom = 21
        nameprocess = get_value_by_key(process_name_dict,addsaderfrom)
        global employee_id
        #print(employee_id)
        datee = datetime.now()
        self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                        VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
        
        self.db.commit()
        
    def Show_ahkam(self):
        '''comboBox_16
        comboBox_2
        comboBox_3'''
        self.comboBox_16.clear()
        self.comboBox_2.clear()
        self.comboBox_3.clear()
        self.comboBox_16.addItem('----')
        self.comboBox_2.addItem('----')
        self.comboBox_3.addItem('----')
        self.cur.execute('''SELECT * FROM cb_ahkam''')
        data = self.cur.fetchall()
        #print(data)
        sorted_data = sorted(data, key=lambda x: x[0])
        for row in sorted_data:
            #print(row[0])
            self.comboBox_16.insertItem(self.comboBox_16.count(), row[1])
            self.comboBox_2.insertItem(self.comboBox_2.count(), row[1])
            self.comboBox_3.insertItem(self.comboBox_3.count(), row[1])

    def del_ahkam(self):
        sf = self.comboBox_16.currentText()
        if sf != "----" :
            reply = QMessageBox.critical(self, "تحذير", "هل تريد بالتأكيد الحذف?", QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.cur.execute('''
                             DELETE FROM cb_ahkam WHERE name=%s 
                             ''',[sf]) 
                self.db.commit()
                self.statusBar().showMessage("تم حذف البند بنجاح")
                self.Show_ahkam()
                delsaderfrom = 28
                nameprocess = get_value_by_key(process_name_dict,delsaderfrom)
                global employee_id
                #print(employee_id)
                datee = datetime.now()
                self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                
                self.db.commit()
            else:
                return

###############################
#### // Handle Login        

    def Handle_Login(self):
        self.statusBar().showMessage("  أدخل اسم المستخدم وكلمة السر الخاصة بك")
        
        self.cur.execute('''
                            SELECT id,name,password FROM employees ''')
        data = self.cur.fetchall()
        # print(data)
        name = self.lineEdit_2.text()
        password=self.lineEdit.text()
        
        for i in data:
            # print(i)
            if i[1]==name and i[2]==password:
                global employee_id
                employee_id = i[0]
                try:
                    self.cur.execute('''SELECT emp_Permisions From employees WHERE id =%s''',[employee_id])
                    onedata= self.cur.fetchone()
                    per = list(onedata[0])
                except Exception as se :
                    self.statusBar().showMessage("يرجى مراجعة قسم المعلوماتية"+str(se))
                # print(per)
                self.Handle_Permission(per)
                loginprocess = 1
                nameprocess = get_value_by_key(process_name_dict,loginprocess)
                datee = datetime.now()
                self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                self.statusBar().showMessage("تم تسجيل الدخول")
                global employee_name
                employee_name = name
                self.Load_User_Settings(employee_name)
                self.db.commit()
                self.Show_All_Operation(employee_id)
                break
            else :
                self.statusBar().showMessage("اسم المستخدم أو كلمة السر خاطئة")
               
    def Hanndle_Reset_Password(self):
        pass
    
    ###############################
    #### // Handle Daily_work
    
    def Show_All_Operation(self,employee_id):
        self.cur.execute('''SELECT id FROM employees WHERE emp_Permisions="abcdefghiklmn"''')
        admins = self.cur.fetchall()
        print(admins)
        for idtemp in admins:
            if idtemp[0] == employee_id :    
                self.tableWidget_5.clear()
                self.tableWidget_5.setColumnCount(4)
                column_name = ['العملية','رقم الموظف','تاريخ العملية','اسم الموظف']
                self.tableWidget_5.setHorizontalHeaderLabels(column_name)
                self.tableWidget_5.show()
                # self.tableWidget_5.insertRow(0)
                self.cur.execute('''
                                SELECT dailymovements.process_type , dailymovements.emp_id , dailymovements.date_process , employees.name 
                                        FROM dailymovements 
                                    JOIN employees ON dailymovements.emp_id = employees.id; ''')

                data = self.cur.fetchall()
                self.tableWidget_5.setRowCount(len(data))
                print(data)
                for row , form in enumerate(data):
                    for col , item in enumerate(form):
                        self.tableWidget_5.setItem(row,col,QTableWidgetItem(str(item)))
                        col+=1
                    # row_pos = self.tableWidget_5.rowCount()
                    # self.tableWidget_5.insertRow(row_pos)
                self.tableWidget_5.resizeColumnsToContents()

                break
            else:
                # self.tableWidget_5.clear()
                # self.tableWidget_5.setColumnCount(4)
                # column_name = ['العملية','رقم الموظف','تاريخ العملية','اسم الموظف']
                # self.tableWidget_5.setHorizontalHeaderLabels(column_name)
                # self.tableWidget_5.show()
                # self.tableWidget_5.insertRow(0)
                self.cur.execute('''Select emp_Permisions FROM employees WHERE id=%s''',[employee_id])
                peremployee = self.cur.fetchone()
                # print(peremployee)
                if 'l' in list(peremployee):
                    self.Show_All_Operation(employee_id)
                else:
                    self.statusBar().showMessage("ليس لديك سماحيات لعرض التقارير")
                    break               

    def Handle_Daily_Work9(self):
        try:
            global employee_id 
            #print(employee_id)
            if self.comboBox_18.currentIndex() == 0 :
                self.statusBar().showMessage("الرجاء تحديد الموظف") 
            else:
                type_process = self.comboBox_9.currentIndex()
                emp_name_index = self.comboBox_18.currentIndex()
                if type_process!= 0:
                    #print(type_process)


                    #self.tableWidget_5.clear()
                    if self.tableWidget_5 is not None:
                        self.tableWidget_5.clear()
                        self.tableWidget_5.setColumnCount(4)
                        column_name = ['العملية','رقم الموظف','تاريخ العملية','اسم الموظف']
                        self.tableWidget_5.setHorizontalHeaderLabels(column_name)
                        self.tableWidget_5.show()
                        # self.tableWidget_5.insertRow(0)
                        self.cur.execute('''SELECT dailymovements.process_type , dailymovements.emp_id , dailymovements.date_process ,
                                        employees.name FROM dailymovements 
                                        JOIN employees ON dailymovements.emp_id = employees.id
                                        WHERE dailymovements.process_type=%s AND employees.id=%s;
                                        ''',[type_process,employee_id])
                        data = self.cur.fetchall()
                        #print(data)
                        self.tableWidget_5.setRowCount(len(data))

                        for row , form in enumerate(data):
                            for col , item in enumerate(form):
                                #print(item)
                                #print(process_name_list[item-1][1])
                                if col == 0 :
                                    self.tableWidget_5.setItem(row,col,QTableWidgetItem(str(process_name_list[item-1][1])))
                                else:
                                    self.tableWidget_5.setItem(row,col,QTableWidgetItem(str(item)))
                                col+=1
                            # row_pos = self.tableWidget_5.rowCount()
                            # self.tableWidget_5.insertRow(row_pos)
                else:
                    #print('done')
                    self.Show_All_Operation(employee_id)

                if emp_name_index!= 0:
                    emp_name = self.comboBox_18.currentText()
                    
                    #print(empid)
                    self.tableWidget_5.clear()
                    self.tableWidget_5.setColumnCount(4)
                    column_name = ['العملية','رقم الموظف','تاريخ العملية','اسم الموظف']
                    self.tableWidget_5.setHorizontalHeaderLabels(column_name)
                    self.tableWidget_5.show()
                    # self.tableWidget_5.insertRow(0)
                    self.cur.execute('''SELECT dailymovements.process_type , dailymovements.emp_id , dailymovements.date_process ,
                                    employees.name FROM dailymovements 
                                    JOIN employees ON dailymovements.emp_id = employees.id
                                    WHERE dailymovements.process_type=%s AND employees.name=%s;
                                    ''',[type_process,emp_name])
                    data = self.cur.fetchall()
                    # print(data)
                    self.tableWidget_5.setRowCount(len(data))

                    for row , form in enumerate(data):
                        for col , item in enumerate(form):
                            if col == 0 :
                                self.tableWidget_5.setItem(row,col,QTableWidgetItem(str(process_name_list[item-1][1])))
                            else:
                                self.tableWidget_5.setItem(row,col,QTableWidgetItem(str(item)))
                            col+=1
                        # row_pos = self.tableWidget_5.rowCount()
                        # self.tableWidget_5.insertRow(row_pos)
        except:
            self.statusBar().showMessage("الجدول غير موجود أو تم حذفه  ") 
  
    def Handle_Daily_Work18(self):
        emp_name = self.comboBox_18.currentText()
        type_process = self.comboBox_9.currentIndex()
        if emp_name!='----':
           
            self.tableWidget_5.clear()
            self.tableWidget_5.setColumnCount(4)
            column_name = ['العملية','رقم الموظف','تاريخ العملية','اسم الموظف']
            self.tableWidget_5.setHorizontalHeaderLabels(column_name)
            self.tableWidget_5.show()
            # self.tableWidget_5.insertRow(0)
            self.cur.execute('''SELECT dailymovements.process_type , dailymovements.emp_id , dailymovements.date_process ,
                              employees.name FROM dailymovements 
                             JOIN employees ON dailymovements.emp_id = employees.id
                             WHERE dailymovements.process_type=%s AND employees.name=%s;
                             ''',[type_process,emp_name])
            data = self.cur.fetchall()
            # print(data)
            self.tableWidget_3.setRowCount(len(data))

            for row , form in enumerate(data):
                for col , item in enumerate(form):
                    if col == 0 :
                        self.tableWidget_5.setItem(row,col,QTableWidgetItem(item))
                    else:
                        self.tableWidget_5.setItem(row,col,QTableWidgetItem(str(item)))
                    col+=1
                # row_pos = self.tableWidget_5.rowCount()
                # self.tableWidget_5.insertRow(row_pos)

########Handle 9--18--33--ComboBox
    
    def Initilaize_Operation(self):
        self.cur.execute('''SELECT name FROM employees''')
        empnames = self.cur.fetchall()
        #print(empnames)
        self.comboBox_18.clear()
        self.comboBox_33.clear()
        self.comboBox_18.addItem("----")
        self.comboBox_33.addItem("----")


        emp_name_list = []
        for i in empnames:
            if i[0] == 'admin':
                pass
            else:
                emp_name_list.append(i[0])
        #print(emp_name_list)
        self.comboBox_18.addItems(emp_name_list)
        self.comboBox_33.addItems(emp_name_list)
      
        # العنصر ذو الرقم 1 في position 0 في الكمبوبوكس
        sorted_names = sorted(process_name_list, key=lambda x: x[0])
        for item in sorted_names:
            self.comboBox_9.addItem(item[1])
      
    def Export_Operation_xls(self):
        from xlsxwriter import Workbook
        # فتح نافذة حفظ الملف
        file_path, _ = QFileDialog.getSaveFileName(
            self, "اختر مكان حفظ الملف", "", "Excel Files (*.xlsx)"
        )
        if not file_path:
            return

        try:
            # إنشاء ملف Excel
            workbook = Workbook(file_path)
            worksheet = workbook.add_worksheet("البيانات")

            # إدراج صورة في الأعلى (اختياري)
            image_path = os.path.join(os.getcwd(), "syrianar3kab.png")
            if os.path.exists(image_path):
                worksheet.insert_image('A1', image_path, {
                    'x_scale': 0.7,
                    'y_scale': 0.7,
                    'object_position': 1
                })
                start_row = 15  # تعديل حسب حجم الصورة
            else:
                start_row = 0

            # تنسيقات الخلايا
            header_format = workbook.add_format({
                'bold': True, 'bg_color': '#D7E4BC',
                'align': 'center', 'border': 1
            })
            cell_format = workbook.add_format({
                'align': 'center', 'border': 1
            })

            # استخراج عدد الصفوف والأعمدة
            column_count = self.tableWidget_5.columnCount()
            row_count = self.tableWidget_5.rowCount()

            # كتابة رؤوس الأعمدة
            for col in range(column_count):
                header_item = self.tableWidget_5.horizontalHeaderItem(col)
                header_text = header_item.text() if header_item else f"عمود_{col + 1}"
                worksheet.write(start_row, col, header_text, header_format)

            # كتابة البيانات
            for row in range(row_count):
                for col in range(column_count):
                    item = self.tableWidget_5.item(row, col)
                    value = item.text() if item else ''
                    worksheet.write(start_row + 1 + row, col, value, cell_format)

            # ضبط عرض الأعمدة تلقائيًا
            for col in range(column_count):
                max_length = len(self.tableWidget_5.horizontalHeaderItem(col).text()) if self.tableWidget_5.horizontalHeaderItem(col) else 10
                for row in range(row_count):
                    item = self.tableWidget_5.item(row, col)
                    if item and item.text():
                        max_length = max(max_length, len(item.text()))
                worksheet.set_column(col, col, max_length + 2)

            # حفظ الملف
            workbook.close()
            self.statusBar().showMessage("✅ تم حفظ الملف بنجاح: " + file_path)

        except Exception as e:
            self.statusBar().showMessage("⚠️ فشل التصدير: " + str(e))

    def Export_Operation_cvs(self):
        # فتح نافذة لحفظ الملف
        file_path, _ = QFileDialog.getSaveFileName(
            self, "اختر مكان الحفظ", "", "CSV Files (*.csv)"
        )
        if not file_path:
            return

        row_count = self.tableWidget_4.rowCount()
        column_count = self.tableWidget_4.columnCount()

        headers = [
            self.tableWidget_4.horizontalHeaderItem(col).text()
            if self.tableWidget_4.horizontalHeaderItem(col)
            else f"عمود_{col + 1}"
            for col in range(column_count)
        ]

        data = []
        for row in range(row_count):
            row_data = [
                self.tableWidget_4.item(row, col).text() if self.tableWidget_4.item(row, col) else ''
                for col in range(column_count)
            ]
            data.append(row_data)

        # إنشاء ملف بصيغة CSV مع تنسيق رمزي
        with open(file_path, 'w', encoding='utf-8-sig', newline='') as f:
            # إدراج نص رمزي يمثل الشعار السوري في بداية الملف
            f.write("🦅 الجمهورية العربية السورية - شعار الدولة\n\n")
            
            # كتابة الجدول
            df = pd.DataFrame(data, columns=headers)
            df.to_csv(f, index=False)

        self.statusBar().showMessage(f"✅ تم حفظ الملف بصيغة CSV مع تنسيق رمزي{file_path}")

#################################################################

    def on_print_button_clicked(self):
        # إنشاء نسخة من الجدول الحالي
        preview_table = QTableWidget(self.tableWidget_5.rowCount(), self.tableWidget_5.columnCount())
        preview_table.setHorizontalHeaderLabels(
            [self.tableWidget_5.horizontalHeaderItem(i).text() for i in range(self.tableWidget_5.columnCount())]
        )
        for row in range(self.tableWidget_5.rowCount()):
            for col in range(self.tableWidget_5.columnCount()):
                item = self.tableWidget_5.item(row, col)
                preview_table.setItem(row, col, QTableWidgetItem(item.text() if item else ""))

        # فتح نافذة المعاينة
        dialog = PrintPreviewDialog(preview_table, self)
        dialog.exec_()
 
    ############//Circle DB//##############
    ####For all users    
    
    def Add_New_Circle(self):
        self.comboBox_5.clear()
        idc = self.comboBox_6.currentText()
        #print(idc)
        self.cur.execute('''
                         SELECT id FROM  department WHERE name=%s''',(idc,))
        dep_idc = self.cur.fetchone()
        #print(dep_idc)
        
        name = self.lineEdit_19.text()
        if name != '----':
            self.cur.execute('''
                            INSERT INTO circle (Circle_Department_id,name)
                            VALUES (%s,%s)
                            ''',(dep_idc[0],name))
            self.db.commit()
        addcircle = 10
        addcirclename = get_value_by_key(process_name_dict,addcircle)
        global employee_id
        #print(employee_id)
        datee = datetime.now()
        self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                         VALUES (%s,%s,%s)''',[addcirclename,employee_id,datee])
        
        self.db.commit()
        # idc_text = Depart_List[idc]
        #idc_text = get_value_by_key(department_choices,dep_idc)
        self.Show_All_Circle_inTable(idc)
        self.lineEdit_19.setText('')
        self.statusBar().showMessage("تم إضافة دائرة بنجاح")
        
    def get_data(self,row,col):
        item = self.tableWidget.item(row, col)
        if item:
            self.statusBar().showMessage(f"القيمة المختارة: {item.text()}")
            # print(f"القيمة المختارة: {item.text()}") 
            self.lineEdit_19.setText(item.text())
            txt = item.text()
            # print(txt) 
        return txt
        
    def Del_Circle(self):
        name_dep = self.comboBox_6.currentText()
        id_dep = get_key_by_value(department_choices,name_dep)
        name_cir = self.lineEdit_19.text()
        # row = self.tableWidget.currentRow()
        # column = self.tableWidget.currentColumn()
        # item = self.tableWidget.item( row, column)
        # r,c,value_del = self.get_data(row,column)
        # value_del = item.text()
        
        if name_cir:
            reply = QMessageBox.critical(self, "تحذير", "هل تريد بالتأكيد الحذف?", QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.cur.execute('''
                             DELETE FROM circle WHERE circle_department_id=%s AND name=%s
                             ''',[id_dep,name_cir]) 
                self.db.commit()
                
                delcircle = 11
                nameprocess = get_value_by_key(process_name_dict,delcircle)
                global employee_id
                #print(employee_id)
                datee = datetime.now()
                self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                
                
                
                self.db.commit()
            else:
                return
    
        self.statusBar().showMessage("تم  حذف دائرة بنجاح")
        self.lineEdit_19.setText('')
        self.Show_All_Circle_inTable(name_dep)
    
    def Show_All_Circle_inTable(self,selected_item):
        self.tableWidget.clear()
        self.tableWidget.setColumnCount(1)
        column_name = ['الاسم']
        self.tableWidget.setHorizontalHeaderLabels(column_name)
        self.tableWidget.show()
        
        # self.tableWidget.insertRow(0)
        if selected_item != '----':
            selected_item_index = get_key_by_value(department_choices,selected_item)
            
            self.cur.execute("""
                        SELECT name FROM circle WHERE circle_department_id = %s
                        """,[selected_item_index])
            data = self.cur.fetchall()
            # print(data)
            self.tableWidget_5.setRowCount(len(data))

            
            for row , form in enumerate(data):
                for col , item in enumerate(form):
                    
                    self.tableWidget.setItem(row,col,QTableWidgetItem(str(item)))
                    col+=1
                # row_pos = self.tableWidget.rowCount()
                # self.tableWidget.insertRow(row_pos)
        else:
            self.statusBar().showMessage("الرجاء التأكد من قيم الادخالات ")
            return

    ############//Permission//################
    # For Admin Only   
    
    def Handle_Permission(self,per):
        if len(per) ==13:
            #print('done')
            self.tabWidget.setEnabled(True)
            self.tabWidget_2.setEnabled(True)
            self.btnentry.setEnabled(True)
            self.btndelete.setEnabled(True)
            self.btnsearch.setEnabled(True)
            self.btndaily.setEnabled(True)
            self.btnreport.setEnabled(True)
            self.btnsettings.setEnabled(True)
            self.pushButton_33.setEnabled(True)

            # self.tabWidget.setTabEnabled(2, True)
           
            self.pushButton_3.setEnabled(True)
            self.pushButton_21.setEnabled(True)
            self.pushButton_3.setEnabled(True)
            self.pushButton_24.setEnabled(True)
            self.pushButton_23.setEnabled(True)
            self.pushButton_14.setEnabled(True)
            self.pushButton_20.setEnabled(True)
            self.pushButton.setEnabled(True)
            self.pushButton_28.setEnabled(True)
            self.pushButton.setEnabled(True)
            self.pushButton_26.setEnabled(True)
            self.pushButton_35.setEnabled(True)
            self.pushButton_25.setEnabled(True)
            self.pushButton_9.setEnabled(True)
            self.pushButton_18.setEnabled(True)
            self.pushButton_10.setEnabled(True)
            self.groupBox_9.setEnabled(True)
            self.pushButton_10.setEnabled(True)
            self.groupBox_10.setEnabled(True)
            self.comboBox_18.setEnabled(True)
            self.pushButton_39.setEnabled(True)
            self.pushButton_40.setEnabled(True)
            self.pushButton_41.setEnabled(True)
            self.comboBox_33.setEnabled(True)
            self.btnPrint.setEnabled(True)
            self.btnExport.setEnabled(True)
            self.pushButton_33.setEnabled(True)
            self.comboBox_8.setEnabled(True)
            self.pushButton_38.setEnabled(True)
            self.pushButton_37.setEnabled(True)

            # self.tabWidget_2.setTabEnabled(0, True)

            self.pushButton_4.setEnabled(True)
            self.pushButton.setEnabled(True)
            self.pushButton_3.setEnabled(True)
            self.pushButton_23.setEnabled(True)
            self.pushButton_10.setEnabled(True)
            self.pushButton_2.setEnabled(True)

            self.comboBox_18.setEnabled(True)
            self.groupBox_10.setEnabled(True)
            self.groupBox_9.setEnabled(True)

            #self.Show_All_Operation()

        else:

            for i in per:
                #add employee
                if i == 'a':
                    print(i)
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(8, True)
                    self.tabWidget_2.setEnabled(True)
                    self.tabWidget_2.setTabEnabled(2, True) 
                    self.btnsettings.setEnabled(True)
                    self.pushButton_3.setEnabled(True)
                    self.pushButton_21.setEnabled(True)
                    
                    
                #delete employee
                elif i == 'b':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(8, True)
                    self.btnsettings.setEnabled(True)
                    self.tabWidget_2.setEnabled(True)
                    self.tabWidget_2.setTabEnabled(3, True)
                    self.pushButton_3.setEnabled(True)
                    self.pushButton_24.setEnabled(True)
                    
                    

                #searchs update employee
                elif i == 'c':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(8, True)
                    self.btnsettings.setEnabled(True)
                    self.tabWidget_2.setEnabled(True)
                    self.tabWidget.setTabEnabled(2, True)
                    self.pushButton_23.setEnabled(True)
                    self.pushButton_14.setEnabled(True)
                    self.pushButton_20.setEnabled(True)
                    
                    
                #add circle
                elif i == 'd':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(8, True)
                    self.btnsettings.setEnabled(True)
                    self.tabWidget_2.setEnabled(True)
                    self.tabWidget.setTabEnabled(1, True)
                    self.pushButton.setEnabled(True)
                    self.pushButton_28.setEnabled(True)
                    

                #delete circle
                elif i == 'e':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(8, True)
                    self.btnsettings.setEnabled(True)
                    self.tabWidget_2.setEnabled(True)
                    self.tabWidget.setTabEnabled(1, True)
                    self.pushButton.setEnabled(True)
                    self.pushButton_26.setEnabled(True)
                    

                #add mail
                elif i == 'f':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(3, True)
                    self.btnentry.setEnabled(True)
                    self.pushButton_35.setEnabled(True)

                #delete mail
                elif i == 'g':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(4, True)
                    self.btndelete.setEnabled(True)
                    self.pushButton_25.setEnabled(True)

                #search and update
                elif i == 'h':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(5, True)
                    self.btnsearch.setEnabled(True)
                    self.pushButton_9.setEnabled(True)
                    self.pushButton_18.setEnabled(True)
                    
                #add to choices
                elif i == 'i':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(8, True)
                    self.btnsettings.setEnabled(True)
                    self.tabWidget_2.setEnabled(True)
                    self.tabWidget.setTabEnabled(4, True)
                    self.pushButton_10.setEnabled(True)
                    self.groupBox_9.setEnabled(True)

                #delete ffrom choices
                elif i == 'k':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(8, True)
                    self.btnsettings.setEnabled(True)
                    self.tabWidget_2.setEnabled(True)
                    self.tabWidget.setTabEnabled(4, True)
                    self.pushButton_10.setEnabled(True)
                    self.groupBox_10.setEnabled(True)

                    
                #dailymovments
                elif i == 'l':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(2, True)
                    self.btndaily.setEnabled(True)
                    self.comboBox_18.setEnabled(True)
                    self.pushButton_39.setEnabled(True)
                    self.pushButton_40.setEnabled(True)
                    self.pushButton_41.setEnabled(True)
                    
                    
                #report
                elif i == 'm':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(7, True)
                    self.btnreport.setEnabled(True)
                    self.comboBox_33.setEnabled(True)
                    self.btnPrint.setEnabled(True)
                    self.btnExport.setEnabled(True)
                    #filter
                elif i == 'n':
                    self.tabWidget.setEnabled(True)
                    self.tabWidget.setTabEnabled(6, True)
                    self.pushButton_33.setEnabled(True)
                    self.comboBox_8.setEnabled(True)
                    self.pushButton_38.setEnabled(True)
                    self.pushButton_37.setEnabled(True)
                #self.Show_All_Operation() 
                # Permision_Dict= {
                    #     'a':'اضافة موظف'
                    #     ,'b':"حذف موظف"
                    #     ,'c':"تعديل موظف"
                    #     ,'d':"اضافة دائرة"
                    #     ,'e':"حذف دائرة"
                    #     ,'f':"اضافة بريد"
                    #     ,'g':"حذف بريد"
                    #     ,'h':"تعديل بريد"
                    #     ,'i':"إضافة إلى الاختيارات"
                    #     ,'k':"حذف من الاختيارات"
                    #     ,'l':"الحركة اليومية"
                    #     ,'m':"التقارير"
                    #     ,'n':"فلترة"
                    # }   
 
   ##############QR Generate##################

    def secure_unique_number(self):
        while True:
            num = secrets.randbelow(10**12)
            if num not in self.used_numbers:
                self.used_numbers.add(num)
                return f"{num:012d}"  # يضمن 12 خانة مع أصفار أمامية

    def generate_qr_data(self):
        public_num_mail = self.lineEdit_7.text()
        enter_date_hidden =datetime.now() 
        date_input_mail = self.dateEdit_12.date().toString("yyyy-MM-dd")
        date_input_mail_str = str(date_input_mail)
        madmoon_mail = self.textEdit_10.toPlainText()
        mostlem_mail = self.comboBox_10.currentText()
        image_path = self.current_image_path
        randomize_Secret_code = self.secure_unique_number()
        # دمج البيانات في سلسلة واحدة
        qr_data = f"{public_num_mail}|{enter_date_hidden}|{date_input_mail_str}|{madmoon_mail}|{mostlem_mail}|{image_path}|{randomize_Secret_code}"
        return qr_data

    def receive_qr(self, pixmap, byte_data):
        resized_pixmap = pixmap.scaled(150, 150, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.qr_label.setPixmap(resized_pixmap)
        self.current_qr_byte = byte_data  # للاستخدام لاحقًا أو الحفظ
    
    def open_qr_window(self):
        try:
            qr_data = self.generate_qr_data()
            if not qr_data.strip():
                QMessageBox.warning(self, "تنبيه", "البيانات المدخلة غير صالحة لتوليد QR.")
                return

            self.qr_window = QRDisplay(qr_data)
            self.qr_window.qr_saved.connect(self.receive_qr)
            self.qr_window.show()

        except Exception as e:
            # print("حدث خطأ أثناء فتح نافذة QR:", e)
            self.statusBar().showMessage(f"حدث خطأ أثناء فتح نافذة QR:{e}" )

    def Print_QR(self):
        # استخراج الصورة من QLabel
        pixmap = self.qr_label.pixmap()
        if pixmap is None or pixmap.isNull():
            # print("❌ لا توجد صورة QR في qr_label.")
            self.statusBar().showMessage("❌ لا توجد صورة QR في qr_label." )

            return

        # إعداد الطابعة
        printer = QPrinter(QPrinter.HighResolution)
        printer.setPageSize(QPrinter.A4)
        printer.setOrientation(QPrinter.Portrait)
        printer.setFullPage(True)

        # فتح نافذة اختيار الطابعة
        dialog = QPrintDialog(printer, self)
        if dialog.exec_() != QPrintDialog.Accepted:
            return

        # إعداد الرسام
        painter = QPainter(printer)

        # تحديد الحجم الحقيقي للطباعة (مثلاً 4×4 سم)
        dpi = printer.resolution()
        cm_to_inches = 1 / 2.54
        target_width_cm = 4
        target_height_cm = 4
        target_width_px = int(target_width_cm * cm_to_inches * dpi)
        target_height_px = int(target_height_cm * cm_to_inches * dpi)

        # حساب موضع الطباعة في منتصف الصفحة
        page_rect = printer.pageRect()
        x = (page_rect.width() - target_width_px) // 2
        y = (page_rect.height() - target_height_px) // 2

        # رسم الصورة بالحجم الحقيقي
        painter.drawPixmap(x, y, target_width_px, target_height_px, pixmap)
        painter.end()

    ##########//Mail_DB//#####################
    
    def Clear_Data_Add(self):
        self.comboBox_7.setCurrentIndex(0)
        self.comboBox_14.setCurrentIndex(0)
        self.comboBox_15.setCurrentIndex(0)
        self.comboBox_11.setCurrentIndex(0)
        self.comboBox_12.setCurrentIndex(0)
        self.comboBox_10.setCurrentIndex(0)
        self.comboBox_13.setCurrentIndex(0)
        self.comboBox_2.setCurrentIndex(0)
        self.lineEdit_7.setText('')
        self.lineEdit_17.setText('')
        self.lineEdit_6.setText('')
        self.lineEdit_20.setText('')
        self.lineEdit_16.setText('')
        self.lineEdit_18.setText('')
        self.lineEdit_24.setText('')
        self.lineEdit_23.setText('')
        self.lineEdit_30.setText('')
        self.lineEdit_31.setText('')
        # self.lineEdit_65.setText('')
        self.lineEdit_32.setText('')
        self.lineEdit_22.setText('')
        self.lineEdit_26.setText('')
        self.lineEdit_27.setText('')
        self.lineEdit_28.setText('')
        self.lineEdit_25.setText('')
        # تصفير الصورة من الـ Label
        self.label_preview.clear()
        self.qr_label.clear()
        # self.label_preview.image = None  # ضروري لمنع إبقاء المرجع للصورة
        self.current_qr_byte = None
        self.current_img_byte = None
        self.current_image_path = None

        self.dateEdit_12.setDate(QDate.currentDate())
        self.dateEdit_3.setDate(QDate.currentDate())
        self.dateEdit_2.setDate(QDate.currentDate())
        self.dateEdit_13.setDate(QDate.currentDate())
        self.dateEdit_14.setDate(QDate.currentDate())

        self.textEdit_10.clear()

    def ClearDataSearch(self):
        self.lineEdit_49.setText('')
        self.lineEdit_9.setText('')
        # self.lineEdit_67.setText('')
        self.lineEdit_49.setText('')
        self.textEdit_11.clear()
        self.dateEdit_16.setDate(QDate.currentDate())
        self.comboBox_21.setCurrentIndex(0)
        self.comboBox_23.setCurrentIndex(0)
        self.comboBox_22.setCurrentIndex(0)
        self.comboBox_19.setCurrentIndex(0)
        self.comboBox_25.setCurrentIndex(0)
        self.comboBox_24.setCurrentIndex(0)
        self.lineEdit_50.setText('')
        self.comboBox_20.setCurrentIndex(0)
        self.lineEdit_40.setText('')
        self.lineEdit_41.setText('')
        self.lineEdit_42.setText('')
        self.lineEdit_43.setText('')
        self.lineEdit_44.setText('')
        self.lineEdit_45.setText('')
        self.lineEdit_46.setText('')
        self.lineEdit_47.setText('')
        self.lineEdit_48.setText('')
        self.comboBox_3.setCurrentIndex(0)
        self.lineEdit_34.setText('')
        self.lineEdit_35.setText('')
        self.lineEdit_37.setText('')
        self.lineEdit_39.setText('')
        self.lineEdit_38.setText('')
        self.dateEdit_15.setDate(QDate.currentDate())
        self.label_preview.clear()

    def Show_All_Mails(self):
    
        self.cur.execute('''SELECT mail.type_mail_id, mail.public_number_mail, mail.mail_date, mail.mail_container, mail.mail_mostlm,
                                    personal_info.name, personal_info.mobile,
                                    personal_info.national, personal_info.mahdar, personal_info.maksam, personal_info.manteka,
                                    personal_info.dareebe, personal_info.mehna, personal_info.segel
                                    FROM mail
                                    LEFT JOIN personal_info ON mail.public_number_mail = personal_info.personal_info_mail

                        ''')
        mail_data = self.cur.fetchall()
        
        # print(mail_data)
        self.tableWidget_3.clear()
        self.tableWidget_3.setColumnCount(14)
        # self.tableWidget_3.setRowCount(0)
        column_name = ['نوع البريد','الرقم العام','تاريخ البريد','مضمون البريد','مستلم البريد',
                       'الاسم الشخصي','الموبايل','الرقم الوطني','المحضر','المقسم','المنطقة','الرقم الضريبي','المهنة','السجل' ]
        self.tableWidget_3.setHorizontalHeaderLabels(column_name)
        self.tableWidget_3.show()
        # self.tableWidget_3.insertRow(0)
        self.tableWidget_3.setRowCount(len(mail_data))
        
        for row , form in enumerate(mail_data):
            for col , item in enumerate(form):
                if col == 0:
                    #print(mail_type_tuble[item-1][1])
                    self.tableWidget_3.setItem(row,col,QTableWidgetItem(mail_type_tuble[item-1][1]))
                else:
                    # self.tableWidget_3.setItem(row,col,QTableWidgetItem(str(item)))
                    value = str(item) if item is not None else "----"
                    self.tableWidget_3.setItem(row, col, QTableWidgetItem(value))

                col+=1
            # row_pos = self.tableWidget_3.rowCount()
            # self.tableWidget_3.insertRow(row_pos)

        self.tableWidget_3.resizeColumnsToContents()

    def Add_mail(self,idmail,public_num_mail):
        enter_date_hidden =datetime.now() 
        date_input_mail = self.dateEdit_12.date().toString("yyyy-MM-dd")
        date_input_mail_str = str(date_input_mail)
        madmoon_mail = self.textEdit_10.toPlainText()
        mostlem_mail = self.comboBox_10.currentText()
        image_path = self.current_image_path
        image_byte = self.current_img_byte
        qr = self.current_qr_byte
        try:
            self.cur.execute("""
            INSERT INTO mail (type_mail_id,public_number_mail,enter_date_hidden,mail_date,mail_container,mail_mostlm,image,image_blob,qr_code) 
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)""",(idmail,public_num_mail,enter_date_hidden,date_input_mail_str,
                                            madmoon_mail,mostlem_mail,image_path,image_byte,qr))
            self.db.commit()
        except Exception as ex:
            self.statusBar().showMessage(f"{ex}حدث خطأ أثناء إدخال البيانات العامة")
            return

    def Add_sader(self,public_num_mail):
        sader_date = self.dateEdit_3.date().toString("yyyy-MM-dd")
        sader_from = self.comboBox_14.currentText()
        sader_to = self.comboBox_15.currentText() 
        try:
            self.cur.execute("""
            INSERT INTO sader (sader_date,sader_from,sader_to,sader_mail) 
            VALUES (%s,%s,%s,%s)""",(sader_date,sader_from,sader_to,public_num_mail))
            self.db.commit()
        except Exception as ex:
                self.statusBar().showMessage(f"{ex}حدث خطأ أثناء إدخال بيانات الصادر")
                return

    def Add_wared(self,public_num_mail):
        wared_date = self.dateEdit_2.date().toString("yyyy-MM-dd")
        wared_from = self.comboBox_11.currentText()
        wared_to = self.comboBox_12.currentText()
        try:
            self.cur.execute("""
                                INSERT INTO wared (wared_date,wared_from,wared_to,wared_mail) 
                                VALUES (%s,%s,%s,%s)""",
                                (wared_date,wared_from,wared_to,public_num_mail))
            self.db.commit()
        except Exception as ex:
                self.statusBar().showMessage(f"{ex}حدث خطأ أثناء إدخال بيانات الصادر")
                return
    
    def Add_ahkam(self,public_num_mail):
        typr_hekem = self.comboBox_2.currentText()
        mod3y_name = self.lineEdit_22.text()
        mod3y_3lih_name = self.lineEdit_26.text()
        karar =self.lineEdit_27.text()
        asas = self.lineEdit_28.text()
        tanfez = self.lineEdit_25.text()
        hekem_date = self.dateEdit_14.date().toString("yyyy-MM-dd")
        try:
            self.cur.execute("""INSERT INTO ahkam (ahkam_type,mh_name,mh3_name,karar,asas,tanfeez,date_ahkam,ahkam_mail) 
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s)""",
                                (typr_hekem,mod3y_name,mod3y_3lih_name,karar,asas,tanfez,hekem_date,public_num_mail))
            self.db.commit()
        except Exception as ex:
            self.statusBar().showMessage(f"{ex}حدث خطأ أثناء إدخال بيانات الصادر")
            return

    def Add_E3trad(self,public_num_mail,national):
        
        
            self.cur.execute("""SELECT idpersonal_info FROM personal_info WHERE national=%s""",[national])
            per_id = self.cur.fetchone()

            mo3tred_name = self.lineEdit_17.text()
            mo3tred_date = self.dateEdit_13.date().toString("yyyy-MM-dd")
            mo3tred_at = self.comboBox_13.currentIndex()
            try:
                self.cur.execute("""
                                INSERT INTO mo3tred (mo3tred_name,mo3tred_date,mo3tred_at,mo3tred_mail,mo3tred_per) 
                                VALUES (%s,%s,%s,%s,%s)""",(mo3tred_name,mo3tred_date,mo3tred_at,public_num_mail,per_id[0]))

                self.db.commit()
            except Exception as ex:
                self.statusBar().showMessage(f"{ex}حدث خطأ أثناء إدخال بيانات الصادر")
                return
    
    def Add_talbat(self,public_num_mail,national):
        self.cur.execute("""SELECT idpersonal_info FROM personal_info WHERE national=%s""",[national])
        talbat_per = self.cur.fetchone()
        try:
            self.cur.execute("""INSERT INTO talbat (talbat_mail,talbat_per) 
                                        VALUES (%s,%s)""",(public_num_mail,talbat_per[0]))
            self.db.commit()
        except Exception as ex:
                self.statusBar().showMessage(f"{ex}حدث خطأ أثناء إدخال بيانات الصادر")
                return

    def Add_aksam(self,public_num_mail):

        fk = self.comboBox_11.currentText()
        tk = self.comboBox_12.currentText()
        if fk != '----':
            fk = self.comboBox_14.currentText()
            tk = self.comboBox_15.currentText()
        try:
            self.cur.execute("""INSERT INTO aksam (from_k,to_k,aksam_mail) 
            VALUES (%s,%s,%s)""",(fk,tk,public_num_mail))
            self.db.commit()
        except Exception as ex:
                self.statusBar().showMessage(f"{ex}حدث خطأ أثناء إدخال بيانات الصادر")
                return

    def Add_Personal(self, public_num_mail):
        
            name = self.lineEdit_6.text().strip()
            akar = self.lineEdit_18.text().strip()
            dareebe = self.lineEdit_30.text().strip()
            mobile_line = self.lineEdit_20.text().strip()
            national = self.lineEdit_16.text().strip()
            manteka = self.lineEdit_23.text().strip()
            maksam = self.lineEdit_24.text().strip()
            mehna = self.lineEdit_31.text().strip()
            segel = self.lineEdit_32.text().strip()

            # # تحقق من صحة رقم الهاتف
            # mobile_line_edit = self.lineEdit_20
            # mobile_text = mobile_line_edit.text().strip()

            # if mobile_text:
            #     if not self.validate_mobile_number(mobile_line_edit):
            #         self.statusBar().showMessage("❌ رقم الهاتف غير صالح. الرجاء إدخال رقم يبدأ بـ 09 أو +963 ويتكون من 10 أرقام")
            #         mobile_line_edit.setStyleSheet("border: 2px solid red;")
            #         mobile_line_edit.setFocus()
            #         return
            #     else:
            #         mobile_line_edit.setStyleSheet("")  # إزالة التلوين إذا كان الرقم صحيح
            #         mobile = mobile_text
            # else:
            #     self.statusBar().showMessage("⚠️ يرجى إدخال رقم الموبايل قبل المتابعة")
            #     mobile_line_edit.setStyleSheet("border: 2px solid orange;")
            #     mobile_line_edit.setFocus()
            #     return

            # if mobile_line:
            #     if not self.validate_mobile_number(self.lineEdit_20):
            #         self.statusBar().showMessage("❌ رقم الهاتف غير صالح")
            #         return
            #     mobile = mobile_line
            # else:
            #     mobile = None
            # استخدم None بدلًا من "" لتوافق قاعدة البيانات
            try:
            # تنفيذ الإدخال
                self.cur.execute("""INSERT INTO personal_info (name, mobile, national, mahdar, maksam, manteka,
                                dareebe, mehna, segel, personal_info_mail) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (name or None, mobile_line or None, national or None, akar or None, maksam or None, manteka or None,
                    dareebe or None, mehna or None, segel or None, public_num_mail or None))

                self.db.commit()
                # print("Rows affected:", self.cur.rowcount)

                self.statusBar().showMessage("✅ تم حفظ البيانات بنجاح")
            except Exception as ex:
                import traceback
                # print("🔍 تتبع الخطأ:\n", traceback.format_exc())
                self.statusBar().showMessage(f"❌ حدث خطأ أثناء إدخال البيانات: {ex}")

    def Add_New_Mail(self):

        idmail = self.comboBox_7.currentIndex()
        self.Handle_comboBox_7(idmail) 

        if idmail != 0 :
            ##sader
            if idmail in (1,2,3,4,5,6):
                public_num_mail = self.lineEdit_7.text()
                self.cur.execute('''SELECT public_number_mail FROM mail''')
                allnumbers = self.cur.fetchall()
                all = []
                for i in allnumbers:
                    all.append(i[0])
                if public_num_mail in all:
                    self.statusBar().showMessage("يرجى إدخال رقم مغاير للرقم العام") 
                    return
                else:
                    self.Add_mail(idmail,public_num_mail)
                    self.Add_sader(public_num_mail)
                    # self.statusBar().showMessage("تم إضافة بريد صادر بنجاح")

                    if self.lineEdit_6.text().strip()!='':
                        self.Add_Personal(public_num_mail)
                        self.statusBar().showMessage("تم إضافة بريد صادر بنجاح")
                        

                    else:
                        self.statusBar().showMessage("  تم إضافة بريد صادر بنجاح لا يوجد معلومات شخصية مدخلة")
                        return
                    
            ##wared
            if idmail in (7,8,9,10) :
                
                public_num_mail = self.lineEdit_7.text()
                self.cur.execute('''SELECT public_number_mail FROM mail''')
                allnumbers = self.cur.fetchall()
                all = []
                for i in allnumbers:
                    all.append(i[0])
                if public_num_mail in all:
                    self.statusBar().showMessage("يرجى إدخال رقم مغاير للرقم العام") 
                    return
                else:
                    self.Add_mail(idmail,public_num_mail)
                    # print("✅ Add_mail executed")
                    self.Add_wared(public_num_mail)
                    # print("✅ Add_sader executed")


                if self.lineEdit_6.text()!="":
                    self.Add_Personal(public_num_mail)
                    # print("✅ Add_Personal executed")
                    self.statusBar().showMessage("تم إضافة بريد وارد بنجاح")

                else:
                    self.statusBar().showMessage("  تم إضافة بريد وارد بنجاح لا يوجد معلومات شخصية مدخلة")
                    return
                             
            ##Ahkam
            if idmail == 11:    
                public_num_mail = self.lineEdit_7.text()
                self.cur.execute('''SELECT public_number_mail FROM mail''')
                allnumbers = self.cur.fetchall()
                all = []
                for i in allnumbers:
                    all.append(i[0])
                if public_num_mail in all:
                    self.statusBar().showMessage("يرجى إدخال رقم مغاير للرقم العام") 
                    return
                else:
                    self.Add_mail(idmail,public_num_mail)
                    self.Add_wared(public_num_mail)
                    self.Add_ahkam(public_num_mail)
                    
                    if self.lineEdit_6.text()!="":
                        self.Add_Personal(public_num_mail)
                        self.statusBar().showMessage("تم إضافة الحكم القضائي بنجاح")
                    else:
                        self.statusBar().showMessage("  تم إضافة الحكم القضائي بنجاح ,لا يوجد معلومات شخصية مدخلة")
                        return
                
            ##keedmali
            if idmail == 12 :
                public_num_mail = self.lineEdit_7.text()
                self.cur.execute('''SELECT public_number_mail FROM mail''')
                allnumbers = self.cur.fetchall()
                all = []
                for i in allnumbers:
                    all.append(i[0])
                if public_num_mail in all:
                    self.statusBar().showMessage("يرجى إدخال رقم مغاير للرقم العام") 
                    return
                else:
                    self.Add_mail(idmail,public_num_mail)
                    self.Add_sader(public_num_mail)
                    national = self.lineEdit_16.text().strip()
                    if self.lineEdit_6.text()!="":
                        self.Add_Personal(public_num_mail)
                        self.statusBar().showMessage("تم إضافة معلومات شخصية بنجاح")
                    else:
                        self.statusBar().showMessage("يرجى إضافة المعلومات الشخصية ")
                        return

                try:
                    self.cur.execute("""SELECT idpersonal_info FROM personal_info WHERE national=%s""",[national])
                    per_id = self.cur.fetchone()
                    self.cur.execute("""
                                        INSERT INTO keedmali (keedmali_mail,keed_per_info) 
                                        VALUES (%s,%s)""",
                                        (public_num_mail,per_id[0]))
                    
                    self.db.commit()
                    self.statusBar().showMessage("تم إضافة قيد مالي / طي ضريبة / دخل مقطوع  بنجاح")
                except Exception as ex:
                    self.statusBar().showMessage(f"{ex}حدث خطأ أثناء إدخال بيانات الصادر")
                    return
           
            ##اعتراضات
            if idmail == 13 :
               
                public_num_mail = self.lineEdit_7.text()
                self.cur.execute('''SELECT public_number_mail FROM mail''')
                allnumbers = self.cur.fetchall()
                all = []
                for i in allnumbers:
                    all.append(i[0])
                if public_num_mail in all:
                    self.statusBar().showMessage("يرجى إدخال رقم مغاير للرقم العام") 
                    return
                else:
                    self.Add_mail(idmail,public_num_mail)
                    self.Add_wared(public_num_mail)


                if self.lineEdit_6.text()!="":
                    
                    self.Add_Personal(public_num_mail)
                    national = self.lineEdit_16.text().strip()
                    if national:
                        self.Add_E3trad(public_num_mail,national)
                        self.statusBar().showMessage("تم إضافة بريد اعتراض بنجاح")
                    else:
                        self.statusBar().showMessage("  تم إضافة بريد اعتراض بنجاح, لا يوجد رقم وطني للشخصية مدخلة")
                        return
                else:
                    self.statusBar().showMessage("  تم إضافة بريد اعتراض بنجاح لا يوجد معلومات شخصية مدخلة")
                    return
                                     
            ## rawateb   
            if idmail == 14 :
                public_num_mail = self.lineEdit_7.text()
                self.cur.execute('''SELECT public_number_mail FROM mail''')
                allnumbers = self.cur.fetchall()
                all = []
                for i in allnumbers:
                    all.append(i[0])
                if public_num_mail in all:
                    self.statusBar().showMessage("يرجى إدخال رقم مغاير للرقم العام") 
                    return
                else:
                    self.Add_mail(idmail,public_num_mail)
                    self.Add_wared(public_num_mail)


                if self.lineEdit_6.text()!="":
                    self.Add_Personal(public_num_mail)
                    self.statusBar().showMessage("تم إضافة بريد وارد بنجاح")

                else:
                    self.statusBar().showMessage("  تم إضافة بريد بيانات رواتب وأجور بنجاح لا يوجد معلومات شخصية مدخلة")
                    return
                

                
                self.statusBar().showMessage("تم إضافة بريد بيانات رواتب وأجور بنجاح")
                 
            #Talbat
            if idmail == 15 or idmail == 16 :
                public_num_mail = self.lineEdit_7.text()
                self.cur.execute('''SELECT public_number_mail FROM mail''')
                allnumbers = self.cur.fetchall()
                all = []
                for i in allnumbers:
                    all.append(i[0])
                if public_num_mail in all:
                    self.statusBar().showMessage("يرجى إدخال رقم مغاير للرقم العام") 
                    return
                else:
                    self.Add_mail(idmail,public_num_mail)
                    self.Add_wared(public_num_mail)


                if self.lineEdit_6.text()!="":
                    self.Add_Personal(public_num_mail)
                    national = self.lineEdit_16.text().strip()
                    if national:
                        self.Add_talbat(public_num_mail,national)
                    self.statusBar().showMessage("تم إضافة طلب بريدي  بنجاح")

                else:
                    self.statusBar().showMessage("  تم إضافة بريد بيانات طلبات , لا يوجد معلومات شخصية مدخلة")
                    return


           
              ##مراسلات أقسام   
            
            ##moraslat aksam
            if idmail == 17 :
                public_num_mail = self.lineEdit_7.text()
                self.cur.execute('''SELECT public_number_mail FROM mail''')
                allnumbers = self.cur.fetchall()
                all = []
                for i in allnumbers:
                    all.append(i[0])
                if public_num_mail in all:
                    self.statusBar().showMessage("يرجى إدخال رقم مغاير للرقم العام") 
                    return
                else:
                    self.Add_mail(idmail,public_num_mail)
                    self.Add_aksam(public_num_mail)
                    self.statusBar().showMessage("تم إضافة مراسلات بين الأقسام بنجاح")

            ##غير ذلك
            if idmail == 18 :
                public_num_mail = self.lineEdit_7.text()
                self.cur.execute('''SELECT public_number_mail FROM mail''')
                allnumbers = self.cur.fetchall()
                all = []
                for i in allnumbers:
                    all.append(i[0])
                if public_num_mail in all:
                    self.statusBar().showMessage("يرجى إدخال رقم مغاير للرقم العام") 
                    return
                else:
                    self.Add_mail(idmail,public_num_mail)
                    self.statusBar().showMessage("تم إضافة بريد عام بنجاح")
                    return
                                
            
            addmail = 4
            nameprocess = get_value_by_key(process_name_dict,addmail)
            global employee_id
            datee = datetime.now()
            self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                         VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
            self.db.commit()
            self.Clear_Data_Add()
            self.Show_All_Mails()
     
        else :
           self.statusBar().showMessage("يرجى إدخال المعلومات بشكل صحيح") 
           self.Add_New_Mail()
           
        self.Show_All_Mails()
        self.Show_All_Operation(employee_id)

    def Del_Mail(self):
        
        public_number = self.lineEdit_33.text().strip()
        # print(type(public_number), public_number)
        if not public_number:
            return
        reply = QMessageBox.critical(self, "تحذير", "هل تريد بالتأكيد الحذف؟", QMessageBox.Yes | QMessageBox.No)
        if reply != QMessageBox.Yes:
            return
        try:
            self.cur.execute('''DELETE FROM mail WHERE public_number_mail=%s;''', [public_number])
            self.db.commit()
        # تسجيل الحركة
            delmail = 5
            nameprocess = get_value_by_key(process_name_dict, delmail)
            global employee_id
            datee = datetime.now()

            self.cur.execute('''INSERT INTO dailymovements (process_type, emp_id, date_process)
                                VALUES (%s, %s, %s)''', [nameprocess, employee_id, datee])
            self.db.commit()

            self.statusBar().showMessage("تم حذف البريد بنجاح")
            self.lineEdit_33.setText("")
            self.Show_All_Operation(employee_id)
            self.Show_All_Mails()

        except Exception as e:
            # QMessageBox.critical(self, "خطأ", f"  _راجع قسم المعلوماتية _حدث خطأ أثناء الحذف:\n{str(e)}")
            self.statusBar().showMessage(f"قسم المعلوماتية _حدث خطأ أثناء الحذف {str(e)}")

            # self.cur.execute('''SELECT type_mail_id FROM mail WHERE public_number_mail=%s;''', [public_number])
            # type_mail = self.cur.fetchone()
            # print("type_mail:", type_mail)

            # per_id = None
            # sader_id = None

            # if type_mail and type_mail[0] in (1, 2, 3, 4, 5, 6):
            #     self.cur.execute('''SELECT idpersonal_info FROM personal_info WHERE personal_info_mail=%s;''', [public_number])
            #     per_id = self.cur.fetchone()
            #     # print("per_id:", per_id)

            #     self.cur.execute('''SELECT id FROM sader WHERE sader_mail=%s;''', [public_number])
            #     sader_id = self.cur.fetchone()
            #     # print("sader_id:", sader_id)

            #     if per_id:
            #         self.cur.execute('''DELETE FROM personal_info WHERE idpersonal_info=%s;''', [per_id[0]])
            #         self.db.commit()

            #     if sader_id:
            #         self.cur.execute('''DELETE FROM sader WHERE id=%s;''', [sader_id[0]])
            #         self.db.commit()

            #     self.cur.execute('''DELETE FROM mail WHERE public_number_mail=%s;''', [public_number])
            #     self.db.commit()
            
            # if type_mail and type_mail[0] in (7,8,9,10):
            #     self.cur.execute('''SELECT idpersonal_info FROM personal_info WHERE personal_info_mail=%s;''', [public_number])
            #     per_id = self.cur.fetchone()
            #     print("per_id:", per_id)

            #     self.cur.execute('''SELECT id FROM wared WHERE wared_mail=%s;''', [public_number])
            #     wared_id = self.cur.fetchone()
            #     print("sader_id:", sader_id)

            #     if per_id:
            #         self.cur.execute('''DELETE FROM personal_info WHERE idpersonal_info=%s;''', [per_id[0]])
            #         self.db.commit()

            #     if wared_id:
            #         self.cur.execute('''DELETE FROM wared WHERE id=%s;''', [wared_id[0]])
            #         self.db.commit()

            #     self.cur.execute('''DELETE FROM mail WHERE public_number_mail=%s;''', [public_number])
            #     self.db.commit()


            # if type_mail and type_mail[0] == 11:
            #     self.cur.execute('''SELECT id FROM ahkam WHERE ahkam_mail=%s;''', [public_number])
            #     ahkam_id = self.cur.fetchone()
            #     print("ahkam_id:", ahkam_id)

            #     self.cur.execute('''SELECT idpersonal_info FROM personal_info WHERE personal_info_mail=%s;''', [public_number])
            #     per_id = self.cur.fetchone()
            #     print("per_id:", per_id)

            #     self.cur.execute('''SELECT id FROM wared WHERE wared_mail=%s;''', [public_number])
            #     wared_id = self.cur.fetchone()
            #     print("sader_id:", wared_id)

            #     if per_id:
            #         self.cur.execute('''DELETE FROM personal_info WHERE idpersonal_info=%s;''', [per_id[0]])
            #         self.db.commit()

            #     if wared_id:
            #         self.cur.execute('''DELETE FROM wared WHERE id=%s;''', [wared_id[0]])
            #         self.db.commit()
            #     if ahkam_id:
            #         self.cur.execute('''DELETE FROM ahkam WHERE id=%s;''', [ahkam_id[0]])
            #         self.db.commit()

            #     self.cur.execute('''DELETE FROM mail WHERE public_number_mail=%s;''', [public_number])
            #     self.db.commit()
            
            # if type_mail and type_mail[0] in (12):
            #     pass
            
            # if type_mail and type_mail[0] in (13):
            #     pass
            
            # if type_mail and type_mail[0] in (14):
            #     pass
            
            # if type_mail and type_mail[0] in (15):
            #     pass
            
            # if type_mail and type_mail[0] in (16):
            #     pass
            
            # if type_mail and type_mail[0] in (17):
            #     pass
         
    def Search_Mail(self):
        self.ClearDataSearch()
    ### For Employee 
        public_number = self.lineEdit_8.text()

        key_word = self.lineEdit_9.text()
        # qr = self.lineEdit_67.text()
        qr = None
        if public_number!='':
            self.cur.execute('''SELECT type_mail_id, enter_date_hidden,mail_date,mail_container,mail_mostlm,
                             image,qr_code FROM mail WHERE public_number_mail = %s;''',[public_number])                             
            all_data = self.cur.fetchall()
            #print(all_data)
            self.lineEdit_49.setText(public_number)
            self.textEdit_11.setPlainText(all_data[0][3])
            date = QDate.fromString(all_data[0][2], "yyyy-MM-dd")
            if date.isValid():
                self.dateEdit_16.setDate(date) 
            self.comboBox_21.setCurrentIndex(all_data[0][0])
            self.cur.execute('''SELECT id FROM cb_mostlem_bareed WHERE name=%s''',[all_data[0][4]])
            idmos = self.cur.fetchone()
            self.comboBox_23.setCurrentIndex(idmos[0])
            #صورة
            if all_data[0][5]!='':
                self.show_image_in_label2(all_data[0][5])
            #qr
            # self.lineEdit_67.setText(all_data[0][6])
            qr = None

            if all_data[0][0] != 0:
                # sader
                if all_data[0][0]==1 or all_data[0][0]==2 or all_data[0][0]==3 or all_data[0][0]==4 or all_data[0][0]==5 or all_data[0][0]==6:
                    self.cur.execute('''SELECT sader_from,sader_to FROM sader WHERE sader_mail = %s;''',[public_number])
                    data = self.cur.fetchone() 
                    # print(data)
                    self.comboBox_22.setCurrentText(data[0])
                    self.comboBox_19.setCurrentText(data[1])


                    self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[public_number])
                    info = self.cur.fetchall()
                    # print(info)
                    self.lineEdit_43.setText(info[0][1])
                    self.lineEdit_44.setText(info[0][2])
                    self.lineEdit_45.setText(info[0][3])

                    self.lineEdit_46.setText(info[0][4])
                    self.lineEdit_47.setText(info[0][5])
                    self.lineEdit_48.setText(info[0][6])

                    self.lineEdit_40.setText(info[0][7])
                    self.lineEdit_41.setText(info[0][8])
                    self.lineEdit_42.setText(info[0][9])
                # #wared
                if all_data[0][0]==7 or all_data[0][0]==8 or all_data[0][0]==9 or all_data[0][0]==10 :
                    self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[public_number])
                    data = self.cur.fetchall() 
                    # print(data)
                    self.comboBox_25.setCurrentText(data[0][0])
                    self.comboBox_24.setCurrentText(data[0][1])
                    self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[public_number])
                    info = self.cur.fetchall()
                    #print(info)
                    self.lineEdit_43.setText(info[0][1])
                    self.lineEdit_44.setText(info[0][2])
                    self.lineEdit_45.setText(info[0][3])

                    self.lineEdit_46.setText(info[0][4])
                    self.lineEdit_47.setText(info[0][5])
                    self.lineEdit_48.setText(info[0][6])

                    self.lineEdit_40.setText(info[0][7])
                    self.lineEdit_41.setText(info[0][8])
                    self.lineEdit_42.setText(info[0][9])
                # #ahkam
                if all_data[0][0] == 11 :
                    
                    self.cur.execute('''SELECT ahkam_type,mh_name,mh3_name,
                                    karar,asas,tanfeez,date_ahkam  FROM ahkam WHERE ahkam_mail = %s;''',[public_number,])           
                    data_ahkam = self.cur.fetchall()
                    #print(data_ahkam)
                    self.comboBox_3.setCurrentText(data_ahkam[0][0])
                    

                    
                    self.lineEdit_34.setText(data_ahkam[0][1])
                    self.lineEdit_35.setText(data_ahkam[0][2])

                    self.lineEdit_37.setText(data_ahkam[0][3])
                    self.lineEdit_39.setText(data_ahkam[0][4])
                    self.lineEdit_38.setText(data_ahkam[0][5])
                    #print(data_ahkam[0][9])
                    datee = QDate.fromString(data_ahkam[0][6], "yyyy-MM-dd")
                    if datee.isValid():
                        self.dateEdit_15.setDate(datee) 

                    self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[public_number])
                    info = self.cur.fetchall()
                    # print(info)
                    self.lineEdit_43.setText(info[0][1])
                    self.lineEdit_44.setText(info[0][2])
                    self.lineEdit_45.setText(info[0][3])

                    self.lineEdit_46.setText(info[0][4])
                    self.lineEdit_47.setText(info[0][5])
                    self.lineEdit_48.setText(info[0][6])

                    self.lineEdit_40.setText(info[0][7])
                    self.lineEdit_41.setText(info[0][8])
                    self.lineEdit_42.setText(info[0][9])
                # Keedmali
                if all_data[0][0] == 12 :
                    self.cur.execute('''SELECT sader_from,sader_to FROM sader WHERE sader_mail = %s;''',[public_number])
                    data = self.cur.fetchall() 
                    # print(data)
                    self.comboBox_22.setCurrentText(data[0][0])
                    self.comboBox_19.setCurrentText(data[0][1])
                    
                    self.cur.execute('''SELECT keed_per_info FROM keedmali WHERE keedmali_mail = %s;''',[public_number])
                    datakeed =  self.cur.fetchone()
                    # print(datakeed)
                    self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[public_number])
                    info = self.cur.fetchall()
                    # print(info)
                    self.lineEdit_43.setText(info[0][1])
                    self.lineEdit_44.setText(info[0][2])
                    self.lineEdit_45.setText(info[0][3])

                    self.lineEdit_46.setText(info[0][4])
                    self.lineEdit_47.setText(info[0][5])
                    self.lineEdit_48.setText(info[0][6])

                    self.lineEdit_40.setText(info[0][7])
                    self.lineEdit_41.setText(info[0][8])
                    self.lineEdit_42.setText(info[0][9])   
                #e3trad
                if all_data[0][0] == 13 :
                    self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[public_number])
                    data = self.cur.fetchall() 
                    #print(data)
                    self.comboBox_25.setCurrentText(data[0][0])
                    self.comboBox_24.setCurrentText(data[0][1])

                    self.cur.execute('''SELECT mo3tred_name,mo3tred_date,mo3tred_at,mo3tred_per FROM mo3tred WHERE mo3tred_mail = %s;''',[public_number])
                    data3 =  self.cur.fetchall()
                    # print(data3)
                    self.lineEdit_50.setText(data3[0][0])
                    date = QDate.fromString(data3[0][1], "yyyy-MM-dd")
                    if date.isValid():
                        self.dateEdit_15.setDate(date)
                    
                    self.comboBox_20.setCurrentIndex(int(data3[0][2]))

                    per = data3[0][3]

                    self.cur.execute('''SELECT name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s AND idpersonal_info=%s''',[public_number,per])
                    info = self.cur.fetchall()
                    #print(info)
                    self.lineEdit_43.setText(info[0][0])
                    self.lineEdit_44.setText(info[0][1])
                    self.lineEdit_45.setText(info[0][2])

                    self.lineEdit_46.setText(info[0][3])
                    self.lineEdit_47.setText(info[0][4])
                    self.lineEdit_48.setText(info[0][5])

                    self.lineEdit_40.setText(info[0][6])
                    self.lineEdit_41.setText(info[0][7])
                    self.lineEdit_42.setText(info[0][8])                   
                #rawateb
                if all_data[0][0] == 14 :
                    self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[public_number])
                    data = self.cur.fetchall() 
                    # self.cur.execute('''SELECT id FROM cb_wared_from WEHRE name=%s''',[])
                    # idf = self.cur.fetchone()
                    # print(data)
                    self.comboBox_25.setCurrentText(data[1][0])
                    self.comboBox_24.setCurrentText(data[1][1])

                    self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[public_number])
                    info = self.cur.fetchall()
                    # print(info)
                    self.lineEdit_43.setText(info[0][1])
                    self.lineEdit_44.setText(info[0][2])
                    self.lineEdit_45.setText(info[0][3])

                    self.lineEdit_46.setText(info[0][4])
                    self.lineEdit_47.setText(info[0][5])
                    self.lineEdit_48.setText(info[0][6])

                    self.lineEdit_40.setText(info[0][7])
                    self.lineEdit_41.setText(info[0][8])
                    self.lineEdit_42.setText(info[0][9])    
                ##مراسلات أقسام 
                if all_data[0][0] == 17 :
                    self.cur.execute('''SELECT sader_from,sader_to FROM sader WHERE sader_mail = %s;''',[public_number])
                    data = self.cur.fetchall() 
                    # print(data)
                    if data:
                        self.comboBox_22.setCurrentText(data[0][0])
                        self.comboBox_19.setCurrentText(data[0][1]) 


                    self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[public_number])
                    data = self.cur.fetchall() 
                    #print(data)
                    if data:
                        self.comboBox_25.setCurrentText(data[0][0])
                        self.comboBox_24.setCurrentText(data[0][1])

                    self.cur.execute('''SELECT from_k,to_k FROM aksam WHERE aksam_mail = %s;''',[public_number])
                    data4 =  self.cur.fetchall()
                   # print(data4)
                ## Talbat
                if all_data[0][0] == 15 or all_data[0][0]== 16:
                    self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[public_number])
                    data = self.cur.fetchall() 
                    # self.cur.execute('''SELECT id FROM cb_wared_from WEHRE name=%s''',[])
                    # idf = self.cur.fetchone()
                    # print(data)
                    self.comboBox_25.setCurrentText(data[0][0])
                    self.comboBox_24.setCurrentText(data[0][1])

                    self.cur.execute('''SELECT talbat_per FROM talbat WHERE talbat_mail=%s;''',[public_number])
                    data_t = self.cur.fetchone()
                    # print(data_t) 
                    self.cur.execute('''SELECT name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s AND idpersonal_info=%s''',[public_number,data_t[0]])
                    info = self.cur.fetchall()
                    #print(info)
                    self.lineEdit_43.setText(info[0][0])
                    self.lineEdit_44.setText(info[0][1])
                    self.lineEdit_45.setText(info[0][2])

                    self.lineEdit_46.setText(info[0][3])
                    self.lineEdit_47.setText(info[0][4])
                    self.lineEdit_48.setText(info[0][5])

                    self.lineEdit_40.setText(info[0][6])
                    self.lineEdit_41.setText(info[0][7])
                    self.lineEdit_42.setText(info[0][8])                                                
        elif key_word!=''  :
            self.cur.execute('''SELECT type_mail_id,public_number_mail,mail_date,mail_container,
                             mail_mostlm,image,qr_code FROM mail WHERE mail_container LIKE "%"%s"%" ''',[key_word])
            dataM = self.cur.fetchall()
            # print(dataM)
            for res in dataM:
                reply = QMessageBox.critical(self, "النتيجة",
                                              f"{res[3]}هل هذه النتيجة صحيحة ?", QMessageBox.Yes | QMessageBox.No)
                if reply == QMessageBox.Yes:
                    
                    self.lineEdit_49.setText(str(res[1]))
                    self.textEdit_11.setPlainText(res[3])
                    date = QDate.fromString(res[2], "yyyy-MM-dd")
                    if date.isValid():
                        self.dateEdit_16.setDate(date) 
                    self.comboBox_21.setCurrentIndex(res[0])
                    self.comboBox_23.setCurrentText(res[4])
                    if res[5]!='':
                        self.show_image_in_label2(res[5])

                    # self.lineEdit_67.setText(res[6])
                    qr = None

                    if res[0]== 1 or  res[0]== 2 or res[0]== 3 or res[0]== 4 or res[0] == 5 or res[0] == 6:
                        self.cur.execute('''SELECT sader_from,sader_to FROM sader WHERE sader_mail = %s;''',[res[1]])
                        data = self.cur.fetchall() 
                        #print(data)
                        self.comboBox_22.setCurrentText(data[0][0])
                        self.comboBox_19.setCurrentText(data[0][1])
                        self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[res[1]])
                        info = self.cur.fetchall()
                        # print(info)
                        self.lineEdit_43.setText(info[0][1])
                        self.lineEdit_44.setText(info[0][2])
                        self.lineEdit_45.setText(info[0][3])

                        self.lineEdit_46.setText(info[0][4])
                        self.lineEdit_47.setText(info[0][5])
                        self.lineEdit_48.setText(info[0][6])

                        self.lineEdit_40.setText(info[0][7])
                        self.lineEdit_41.setText(info[0][8])
                        self.lineEdit_42.setText(info[0][9])

                    if res[0]==7 or res[0]==8 or res[0]==9 or res[0]==10 :
                        self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[res[1]])
                        data = self.cur.fetchall() 
                        #print(data)
                        self.comboBox_25.setCurrentText(data[0][0])
                        self.comboBox_24.setCurrentText(data[0][1])
                        self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[res[1]])
                        info = self.cur.fetchall()
                        # print(info)
                        self.lineEdit_43.setText(info[0][1])
                        self.lineEdit_44.setText(info[0][2])
                        self.lineEdit_45.setText(info[0][3])

                        self.lineEdit_46.setText(info[0][4])
                        self.lineEdit_47.setText(info[0][5])
                        self.lineEdit_48.setText(info[0][6])

                        self.lineEdit_40.setText(info[0][7])
                        self.lineEdit_41.setText(info[0][8])
                        self.lineEdit_42.setText(info[0][9])
                    
                    if res[0] == 11 :
                        self.cur.execute('''SELECT ahkam_type,mh_name,mh3_name,
                                        karar,asas,tanfeez,date_ahkam  FROM ahkam WHERE ahkam_mail = %s;''',[res[1],])           
                        data_ahkam = self.cur.fetchall()
                        #print(data_ahkam)
                        self.comboBox_3.setCurrentText(data_ahkam[0][0])
                        
                        self.lineEdit_34.setText(data_ahkam[0][1])
                        self.lineEdit_35.setText(data_ahkam[0][2])

                        self.lineEdit_37.setText(data_ahkam[0][3])
                        self.lineEdit_39.setText(data_ahkam[0][4])
                        self.lineEdit_38.setText(data_ahkam[0][5])

                        datee = QDate.fromString(data_ahkam[0][6], "yyyy-MM-dd")
                        if datee.isValid():
                            self.dateEdit_15.setDate(datee) 
                        
                        self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[res[1]])
                        info = self.cur.fetchall()
                        # print(info)
                        self.lineEdit_43.setText(info[0][1])
                        self.lineEdit_44.setText(info[0][2])
                        self.lineEdit_45.setText(info[0][3])

                        self.lineEdit_46.setText(info[0][4])
                        self.lineEdit_47.setText(info[0][5])
                        self.lineEdit_48.setText(info[0][6])

                        self.lineEdit_40.setText(info[0][7])
                        self.lineEdit_41.setText(info[0][8])
                        self.lineEdit_42.setText(info[0][9])
                   
                    if res[0]  == 12 :
                        self.cur.execute('''SELECT sader_from,sader_to FROM sader WHERE sader_mail = %s;''',[res[1] ])
                        data = self.cur.fetchall() 
                        #print(data)
                        self.comboBox_22.setCurrentText(data[0][0])
                        self.comboBox_19.setCurrentText(data[0][1])
                        self.cur.execute('''SELECT keed_per_info FROM keedmali WHERE keedmali_mail = %s;''',[res[1] ])
                        datakeed =  self.cur.fetchone()
                        #print(datakeed)

                        self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[res[1]])
                        info = self.cur.fetchall()
                        # print(info)
                        self.lineEdit_43.setText(info[0][1])
                        self.lineEdit_44.setText(info[0][2])
                        self.lineEdit_45.setText(info[0][3])

                        self.lineEdit_46.setText(info[0][4])
                        self.lineEdit_47.setText(info[0][5])
                        self.lineEdit_48.setText(info[0][6])

                        self.lineEdit_40.setText(info[0][7])
                        self.lineEdit_41.setText(info[0][8])
                        self.lineEdit_42.setText(info[0][9])
                    
                    if res[0]  == 13 :
                        self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[res[1]])
                        data = self.cur.fetchall() 
                        #print(data)
                        self.comboBox_25.setCurrentText(data[0][0])
                        self.comboBox_24.setCurrentText(data[0][1])
                        self.cur.execute('''SELECT mo3tred_name,mo3tred_date,mo3tred_at,mo3tred_per FROM mo3tred WHERE mo3tred_mail = %s;''',[res[1]])
                        data3 =  self.cur.fetchall()
                        # print(data3)
                        self.lineEdit_50.setText(data3[0][0])
                        date = QDate.fromString(data3[0][1], "yyyy-MM-dd")
                        if date.isValid():
                            self.dateEdit_15.setDate(date)
                        
                        self.comboBox_20.setCurrentIndex(int(data3[0][2]))

                        per = data3[0][3]

                        self.cur.execute('''SELECT name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                        WHERE personal_info_mail=%s AND idpersonal_info=%s''',[res[1],per])
                        info = self.cur.fetchall()
                        #print(info)
                        self.lineEdit_43.setText(info[0][0])
                        self.lineEdit_44.setText(info[0][1])
                        self.lineEdit_45.setText(info[0][2])

                        self.lineEdit_46.setText(info[0][3])
                        self.lineEdit_47.setText(info[0][4])
                        self.lineEdit_48.setText(info[0][5])

                        self.lineEdit_40.setText(info[0][6])
                        self.lineEdit_41.setText(info[0][7])
                        self.lineEdit_42.setText(info[0][8])
                                            
                    if res[0]  == 14 :
                        self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[res[1]])
                        data = self.cur.fetchall() 
                        # self.cur.execute('''SELECT id FROM cb_wared_from WEHRE name=%s''',[])
                        # idf = self.cur.fetchone()
                        # print(data)
                        self.comboBox_25.setCurrentText(data[1][0])
                        self.comboBox_24.setCurrentText(data[1][1])

                        self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                        WHERE personal_info_mail=%s''',[res[1]])
                        info = self.cur.fetchall()
                        # print(info)
                        self.lineEdit_43.setText(info[0][1])
                        self.lineEdit_44.setText(info[0][2])
                        self.lineEdit_45.setText(info[0][3])

                        self.lineEdit_46.setText(info[0][4])
                        self.lineEdit_47.setText(info[0][5])
                        self.lineEdit_48.setText(info[0][6])

                        self.lineEdit_40.setText(info[0][7])
                        self.lineEdit_41.setText(info[0][8])
                        self.lineEdit_42.setText(info[0][9])  
                    
                    if res[0] == 17 :
                        self.cur.execute('''SELECT sader_from,sader_to FROM sader WHERE sader_mail = %s;''',[res[1]])
                        data = self.cur.fetchall() 
                        # print(data)
                        if data:
                            self.comboBox_22.setCurrentText(data[0][0])
                            self.comboBox_19.setCurrentText(data[0][1]) 


                        self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[res[1]])
                        data = self.cur.fetchall() 
                        #print(data)
                        if data:
                            self.comboBox_25.setCurrentText(data[0][0])
                            self.comboBox_24.setCurrentText(data[0][1])
                    
                    if res[0]  == 15 or res[0] == 16:
                        self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[res[1]])
                        data = self.cur.fetchall() 
                        # self.cur.execute('''SELECT id FROM cb_wared_from WEHRE name=%s''',[])
                        # idf = self.cur.fetchone()
                        # print(data)
                        self.comboBox_25.setCurrentText(data[1][0])
                        self.comboBox_24.setCurrentText(data[1][1])
                        self.cur.execute('''SELECT talbat_per FROM talbat WHERE talbat_mail=%s;''',[res[1]])
                        data_t = self.cur.fetchone()
                        # print(data_t) 
                        self.cur.execute('''SELECT name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                        WHERE personal_info_mail=%s ''',[res[1]])
                        info = self.cur.fetchall()
                        #print(info)
                        self.lineEdit_43.setText(info[0][0])
                        self.lineEdit_44.setText(info[0][1])
                        self.lineEdit_45.setText(info[0][2])

                        self.lineEdit_46.setText(info[0][3])
                        self.lineEdit_47.setText(info[0][4])
                        self.lineEdit_48.setText(info[0][5])

                        self.lineEdit_40.setText(info[0][6])
                        self.lineEdit_41.setText(info[0][7])
                        self.lineEdit_42.setText(info[0][8]) 

                    break
                    
                else:
                    pass
        elif qr !='':
            self.cur.execute('''SELECT type_mail_id, enter_date_hidden,mail_date,mail_container,mail_mostlm,
                             image,public_number_mail FROM mail WHERE qr_code = %s;''',[qr])                             
            all_data = self.cur.fetchall()
            #print(all_data)
            self.lineEdit_49.setText(str(all_data[0][6]))
            self.textEdit_11.setPlainText(all_data[0][3])
            date = QDate.fromString(all_data[0][2], "yyyy-MM-dd")
            if date.isValid():
                self.dateEdit_16.setDate(date) 
            self.comboBox_21.setCurrentIndex(all_data[0][0])
            self.cur.execute('''SELECT id FROM cb_mostlem_bareed WHERE name=%s''',[all_data[0][4]])
            idmos = self.cur.fetchone()
            self.comboBox_23.setCurrentIndex(idmos[0])
            #صورة
            if all_data[0][5]!='':
                self.show_image_in_label2(all_data[0][5])
            

            if all_data[0][0] != 0:
                # sader
                if all_data[0][0]==1 or all_data[0][0]==2 or all_data[0][0]==3 or all_data[0][0]==4 or all_data[0][0]==5 or all_data[0][0]==6:
                    self.cur.execute('''SELECT sader_from,sader_to FROM sader WHERE sader_mail = %s;''',[all_data[0][6]])
                    data = self.cur.fetchone() 
                    # print(data)
                    self.comboBox_22.setCurrentText(data[0])
                    self.comboBox_19.setCurrentText(data[1])


                    self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[all_data[0][6]])
                    info = self.cur.fetchall()
                    # print(info)
                    self.lineEdit_43.setText(info[0][1])
                    self.lineEdit_44.setText(info[0][2])
                    self.lineEdit_45.setText(info[0][3])

                    self.lineEdit_46.setText(info[0][4])
                    self.lineEdit_47.setText(info[0][5])
                    self.lineEdit_48.setText(info[0][6])

                    self.lineEdit_40.setText(info[0][7])
                    self.lineEdit_41.setText(info[0][8])
                    self.lineEdit_42.setText(info[0][9])
                # #wared
                if all_data[0][0]==7 or all_data[0][0]==8 or all_data[0][0]==9 or all_data[0][0]==10 :
                    self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[all_data[0][6]])
                    data = self.cur.fetchall() 
                    # print(data)
                    self.comboBox_25.setCurrentText(data[0][0])
                    self.comboBox_24.setCurrentText(data[0][1])
                    self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[all_data[0][6]])
                    info = self.cur.fetchall()
                    #print(info)
                    self.lineEdit_43.setText(info[0][1])
                    self.lineEdit_44.setText(info[0][2])
                    self.lineEdit_45.setText(info[0][3])

                    self.lineEdit_46.setText(info[0][4])
                    self.lineEdit_47.setText(info[0][5])
                    self.lineEdit_48.setText(info[0][6])

                    self.lineEdit_40.setText(info[0][7])
                    self.lineEdit_41.setText(info[0][8])
                    self.lineEdit_42.setText(info[0][9])
                # #ahkam
                if all_data[0][0] == 11 :
                    
                    self.cur.execute('''SELECT ahkam_type,mh_name,mh3_name,
                                    karar,asas,tanfeez,date_ahkam  FROM ahkam WHERE ahkam_mail = %s;''',[all_data[0][6],])           
                    data_ahkam = self.cur.fetchall()
                    #print(data_ahkam)
                    self.comboBox_3.setCurrentText(data_ahkam[0][0])
 
                    self.lineEdit_34.setText(data_ahkam[0][1])
                    self.lineEdit_35.setText(data_ahkam[0][2])

                    self.lineEdit_37.setText(data_ahkam[0][3])
                    self.lineEdit_39.setText(data_ahkam[0][4])
                    self.lineEdit_38.setText(data_ahkam[0][5])
                    #print(data_ahkam[0][9])
                    datee = QDate.fromString(data_ahkam[0][6], "yyyy-MM-dd")
                    if datee.isValid():
                        self.dateEdit_15.setDate(datee) 

                    self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[all_data[0][6]])
                    info = self.cur.fetchall()
                    # print(info)
                    self.lineEdit_43.setText(info[0][1])
                    self.lineEdit_44.setText(info[0][2])
                    self.lineEdit_45.setText(info[0][3])

                    self.lineEdit_46.setText(info[0][4])
                    self.lineEdit_47.setText(info[0][5])
                    self.lineEdit_48.setText(info[0][6])

                    self.lineEdit_40.setText(info[0][7])
                    self.lineEdit_41.setText(info[0][8])
                    self.lineEdit_42.setText(info[0][9])
                # Keedmali
                if all_data[0][0] == 12 :
                    self.cur.execute('''SELECT sader_from,sader_to FROM sader WHERE sader_mail = %s;''',[all_data[0][6]])
                    data = self.cur.fetchall() 
                    # print(data)
                    self.comboBox_22.setCurrentText(data[0][0])
                    self.comboBox_19.setCurrentText(data[0][1])
                    
                    self.cur.execute('''SELECT keed_per_info FROM keedmali WHERE keedmali_mail = %s;''',[all_data[0][6]])
                    datakeed =  self.cur.fetchone()
                    # print(datakeed)
                    self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[all_data[0][6]])
                    info = self.cur.fetchall()
                    # print(info)
                    self.lineEdit_43.setText(info[0][1])
                    self.lineEdit_44.setText(info[0][2])
                    self.lineEdit_45.setText(info[0][3])

                    self.lineEdit_46.setText(info[0][4])
                    self.lineEdit_47.setText(info[0][5])
                    self.lineEdit_48.setText(info[0][6])

                    self.lineEdit_40.setText(info[0][7])
                    self.lineEdit_41.setText(info[0][8])
                    self.lineEdit_42.setText(info[0][9])   
                #e3trad
                if all_data[0][0] == 13 :
                    self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[all_data[0][6]])
                    data = self.cur.fetchall() 
                    #print(data)
                    self.comboBox_25.setCurrentText(data[0][0])
                    self.comboBox_24.setCurrentText(data[0][1])

                    self.cur.execute('''SELECT mo3tred_name,mo3tred_date,mo3tred_at,mo3tred_per FROM mo3tred WHERE mo3tred_mail = %s;''',[all_data[0][6]])
                    data3 =  self.cur.fetchall()
                    # print(data3)
                    self.lineEdit_50.setText(data3[0][0])
                    date = QDate.fromString(data3[0][1], "yyyy-MM-dd")
                    if date.isValid():
                        self.dateEdit_15.setDate(date)
                    
                    self.comboBox_20.setCurrentIndex(int(data3[0][2]))

                    per = data3[0][3]

                    self.cur.execute('''SELECT name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s AND idpersonal_info=%s''',[all_data[0][6],per])
                    info = self.cur.fetchall()
                    #print(info)
                    self.lineEdit_43.setText(info[0][0])
                    self.lineEdit_44.setText(info[0][1])
                    self.lineEdit_45.setText(info[0][2])

                    self.lineEdit_46.setText(info[0][3])
                    self.lineEdit_47.setText(info[0][4])
                    self.lineEdit_48.setText(info[0][5])

                    self.lineEdit_40.setText(info[0][6])
                    self.lineEdit_41.setText(info[0][7])
                    self.lineEdit_42.setText(info[0][8])                   
                #rawateb
                if all_data[0][0] == 14 :
                    self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[all_data[0][6]])
                    data = self.cur.fetchall() 
                    # self.cur.execute('''SELECT id FROM cb_wared_from WEHRE name=%s''',[])
                    # idf = self.cur.fetchone()
                    # print(data)
                    self.comboBox_25.setCurrentText(data[1][0])
                    self.comboBox_24.setCurrentText(data[1][1])

                    self.cur.execute('''SELECT 'idpersonal_info',name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s''',[all_data[0][6]])
                    info = self.cur.fetchall()
                    # print(info)
                    self.lineEdit_43.setText(info[0][1])
                    self.lineEdit_44.setText(info[0][2])
                    self.lineEdit_45.setText(info[0][3])

                    self.lineEdit_46.setText(info[0][4])
                    self.lineEdit_47.setText(info[0][5])
                    self.lineEdit_48.setText(info[0][6])

                    self.lineEdit_40.setText(info[0][7])
                    self.lineEdit_41.setText(info[0][8])
                    self.lineEdit_42.setText(info[0][9])    
                ##مراسلات أقسام 
                if all_data[0][0] == 17 :
                    self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[all_data[0][6]])
                    data = self.cur.fetchall() 
                    # self.cur.execute('''SELECT id FROM cb_wared_from WEHRE name=%s''',[])
                    # idf = self.cur.fetchone()
                    # print(data)
                    self.comboBox_25.setCurrentText(data[1][0])
                    self.comboBox_24.setCurrentText(data[1][1])
                    self.cur.execute('''SELECT sader_from,sader_to FROM sader WHERE sader_mail = %s;''',[all_data[0][6]])
                    data = self.cur.fetchall() 
                    # print(data)
                    if data:
                        self.comboBox_22.setCurrentText(data[0][0])
                        self.comboBox_19.setCurrentText(data[0][1]) 


                    self.cur.execute('''SELECT wared_from,wared_to FROM wared WHERE wared_mail = %s;''',[all_data[0][6]])
                    data = self.cur.fetchall() 
                    #print(data)
                    if data:
                        self.comboBox_25.setCurrentText(data[0][0])
                        self.comboBox_24.setCurrentText(data[0][1])

                    self.cur.execute('''SELECT from_k,to_k FROM aksam WHERE aksam_mail = %s;''',[all_data[0][6]])
                    data4 =  self.cur.fetchall()
                   # print(data4)
                ## Talbat
                if all_data[0][0] == 15 or all_data[0][0]== 16:
                    self.cur.execute('''SELECT talbat_per FROM talbat WHERE talbat_mail=%s;''',[all_data[0][6]])
                    data_t = self.cur.fetchone()
                    # print(data_t) 
                    self.cur.execute('''SELECT name,mobile,national,mahdar,maksam,manteka,dareebe,mehna,segel FROM personal_info
                                     WHERE personal_info_mail=%s AND idpersonal_info=%s''',[all_data[0][6],data_t[0]])
                    info = self.cur.fetchall()
                    #print(info)
                    self.lineEdit_43.setText(info[0][0])
                    self.lineEdit_44.setText(info[0][1])
                    self.lineEdit_45.setText(info[0][2])

                    self.lineEdit_46.setText(info[0][3])
                    self.lineEdit_47.setText(info[0][4])
                    self.lineEdit_48.setText(info[0][5])

                    self.lineEdit_40.setText(info[0][6])
                    self.lineEdit_41.setText(info[0][7])
                    self.lineEdit_42.setText(info[0][8])               

##############QR Search##################

    def show_qr_image(self,qr_bytes, label: QLabel):
        pixmap = QPixmap()
        if pixmap.loadFromData(qr_bytes):
            label.setPixmap(pixmap)
            # print("✅ تم عرض الصورة بنجاح")
        else:
            label.setText("❌ فشل تحميل الصورة")

    def search_database(self,public_number):
        self.cur.execute("""SELECT type_mail_id,enter_date_hidden,mail_date,mail_container,mail_mostlm,image,qr_code FROM mail WHERE public_number_mail=%s""",[public_number])
        info = self.cur.fetchone()
        # print(info[0])
        self.ClearDataSearch()
        self.lineEdit_49.setText(public_number)
        self.textEdit_11.setPlainText(info[3])
        date = QDate.fromString(info[2], "yyyy-MM-dd")
        if date.isValid():
            self.dateEdit_16.setDate(date) 
        self.comboBox_21.setCurrentIndex(info[0])
        self.cur.execute('''SELECT id FROM cb_mostlem_bareed WHERE name=%s''',[info[4]])
        idmos = self.cur.fetchone()
        self.comboBox_23.setCurrentIndex(idmos[0])
        
        #صورة
        if info[5]!='':
            self.show_image_in_label2(info[5])
        else :
            self.statusBar().showMessage("Please See the IT section.")
        if info[6] != None:
            self.show_qr_image(info[6],self.qr_label_2)
        
        return info[0]

    def Search_QR(self):
        if hasattr(self, 'qr_window') and self.qr_window.isVisible():
            QMessageBox.information(self, "تنبيه", "نافذة المسح مفتوحة بالفعل.")
            return

        self.qr_window = QRScannerWindow()
        self.qr_window.qr_saved.connect(self.receive_qr_scan)
        self.qr_window.show()

    def receive_qr_scan(self, data):
        self.current_qr_data = data
          # للاستخدام لاحقًا أو الحفظ
        self.search_by_qr(self.current_qr_data)

    def search_by_qr(self, qr_data):
        # print("تم استقبال QR:", qr_data)
        first_value = qr_data.split('|')[0]
        # first_value = result.split('|')[0]
        result = self.search_database(first_value)
        if result:
            # print(result)
            pass
        else:
            QMessageBox.information(self, "نتيجة البحث", "لم يتم العثور على بيانات مطابقة.")

############Update mail############
    
    def Update_mail(self):
        
        pn = self.lineEdit_49.text()  
        mad = self.textEdit_11.toPlainText()    
        eh = datetime.now()            
        da = self.dateEdit_16.date().toString("yyyy-MM-dd")
        mail_type = self.comboBox_21.currentIndex()
        mos = self.comboBox_23.currentText()
        img = self.current_image_path 

        sfr = self.comboBox_22.currentText()
        sto = self.comboBox_19.currentText()

        wfr = self.comboBox_25.currentIndex()
        wto = self.comboBox_24.currentIndex()

        m3name = self.lineEdit_50.text()
        m3date = self.dateEdit_15.date().toString("yyyy-MM-dd")   
        m3at = self.comboBox_20.currentIndex()

        name = self.lineEdit_43.text()
        mobile_line = self.lineEdit_44
        if not self.validate_mobile_number(mobile_line):
            self.statusBar().showMessage("❌ رقم الهاتف غير صالح")
            return  # أوقف العملية بأمان
        mobile = mobile_line.text().strip()
        
        national = self.lineEdit_45.text()

        mahder =self.lineEdit_46.text()
        maksam =self.lineEdit_47.text()
        manteka =self.lineEdit_48.text()

        dareebe =  self.lineEdit_40.text()
        mehna =  self.lineEdit_41.text()
        segel =  self.lineEdit_42.text()

        type_hekem = self.comboBox_3.currentIndex()   
        mh = self.lineEdit_34.text()
        mha = self.lineEdit_35.text()
        s37 = self.lineEdit_37.text()
        s39 = self.lineEdit_39.text()
        s38 = self.lineEdit_38.text()

                 
        if mail_type== 1 or  mail_type== 2 or mail_type== 3 or mail_type== 4 or mail_type== 5 or mail_type== 6:
            try:        
                self.cur.execute('''UPDATE mail SET type_mail_id=%s,enter_date_hidden=%s,mail_date=%s,
                                mail_container=%s,mail_mostlm=%s,image=%s WHERE public_number_mail=%s''',[mail_type,eh,da,mad,mos,img,pn])
                self.db.commit()            
                
                self.cur.execute('''UPDATE sader SET sader_from=%s,sader_to=%s WHERE sader_mail=%s''',[sfr,sto,pn])
                self.db.commit()
                        
                self.cur.execute('''UPDATE personal_info SET name=%s,mobile=%s,national=%s,mahdar=%s,maksam=%s,manteka=%s,dareebe=%s,mehna=%s,segel=%s
                                WHERE personal_info_mail=%s''',[name,mobile,national,mahder,maksam,manteka,dareebe,mehna,segel,pn])
                self.db.commit()

                self.statusBar().showMessage("تم تحديث بريد صادر بنجاح")
            except Exception as ex:
                self.statusBar().showMessage(f"{ex} عذرا ,, يوجد خطأ في المدخلات ..يرجى التأكد من المعلومات")

        if   mail_type== 7 or mail_type== 8 or mail_type== 9 or mail_type== 10 :
            try:
                self.cur.execute('''UPDATE mail SET type_mail_id=%s,enter_date_hidden=%s,mail_date=%s,
                                mail_container=%s,mail_mostlm=%s,image=%s WHERE public_number_mail=%s''',[mail_type,eh,da,mad,mos,img,pn])
                self.db.commit()

                self.cur.execute('''UPDATE wared SET wared_from=%s,wared_to=%s WHERE wared_mail=%s''',[wfr,wto,pn])
                self.db.commit()

                self.cur.execute('''UPDATE personal_info SET name=%s,mobile=%s,national=%s,mahdar=%s,maksam=%s,manteka=%s,dareebe=%s,mehna=%s,segel=%s
                                    WHERE personal_info_mail=%s''',[name,mobile,national,mahder,maksam,manteka,dareebe,mehna,segel,pn])
                self.db.commit()

                self.statusBar().showMessage("تم تحديث بريد وارد بنجاح")
            except Exception as ex:
                self.statusBar().showMessage(f"{ex} عذرا ,, يوجد خطأ في المدخلات ..يرجى التأكد من المعلومات")

        if mail_type== 11 :
            
            try:
                self.cur.execute('''UPDATE mail SET type_mail_id=%s,enter_date_hidden=%s,mail_date=%s,
                                mail_container=%s,mail_mostlm=%s,image=%s WHERE public_number_mail=%s''',[mail_type,eh,da,mad,mos,img,pn])
                self.db.commit()
                self.cur.execute('''UPDATE ahkam SET ahkam_type=%s,mh_name=%s,mh3_name=%s
                                ,karar=%s,asas=%s,tanfeez=%s,date_ahkam=%s
                                WHERE ahkam_mail=%s''',[type_hekem,mh,mha,s37,s39,s38,m3date,pn])
                self.db.commit()
                self.cur.execute('''UPDATE personal_info SET name=%s,mobile=%s,national=%s,mahdar=%s,maksam=%s,manteka=%s,dareebe=%s,mehna=%s,segel=%s
                                    WHERE personal_info_mail=%s''',[name,mobile,national,mahder,maksam,manteka,dareebe,mehna,segel,pn])
                self.db.commit()
                self.statusBar().showMessage("تم تحديث بريد حكم قضائي بنجاح")
            except Exception as ex:
                self.statusBar().showMessage(f"{ex} عذرا ,, يوجد خطأ في المدخلات ..يرجى التأكد من المعلومات")
            
        if mail_type == 12 :

            try:
                self.cur.execute('''UPDATE mail SET type_mail_id=%s,enter_date_hidden=%s,mail_date=%s,
                                mail_container=%s,mail_mostlm=%s,image=%s WHERE public_number_mail=%s''',[mail_type,eh,da,mad,mos,img,pn])
                self.db.commit()
                self.cur.execute('''UPDATE sader SET sader_from=%s,sader_to=%s WHERE sader_mail=%s''',[sfr,sto,pn])
                self.db.commit()
                self.cur.execute('''UPDATE personal_info SET name=%s,mobile=%s,national=%s,mahdar=%s,maksam=%s,manteka=%s,dareebe=%s,mehna=%s,segel=%s
                                        WHERE personal_info_mail=%s''',[name,mobile,national,mahder,maksam,manteka,dareebe,mehna,segel,pn])
                self.db.commit()
                self.statusBar().showMessage("تم تحديث بريد قيد مالي بنجاح")
            except Exception as ex:
                self.statusBar().showMessage(f"{ex} عذرا ,, يوجد خطأ في المدخلات ..يرجى التأكد من المعلومات")
            
        if mail_type == 13 :
            try:
                self.cur.execute('''UPDATE mail SET type_mail_id=%s,enter_date_hidden=%s,mail_date=%s,
                                mail_container=%s,mail_mostlm=%s,image=%s WHERE public_number_mail=%s''',[mail_type,eh,da,mad,mos,img,pn])
                self.db.commit()
                self.cur.execute('''UPDATE wared SET wared_from=%s,wared_to=%s WHERE wared_mail=%s''',[wfr,wto,pn])
                self.db.commit()
                self.cur.execute('''UPDATE mo3tred SET mo3tred_name=%s,mo3tred_date=%s,mo3tred_at=%s
                                WHERE mo3tred_mail=%s''',[m3name,m3date,m3at,pn])
                self.db.commit()
                self.cur.execute('''UPDATE personal_info SET name=%s,mobile=%s,national=%s,mahdar=%s,maksam=%s,manteka=%s,dareebe=%s,mehna=%s,segel=%s
                                    WHERE personal_info_mail=%s''',[name,mobile,national,mahder,maksam,manteka,dareebe,mehna,segel,pn])
                self.db.commit()
                self.statusBar().showMessage("تم تحديث بريد اعتراض بنجاح")
            except Exception as ex:
                self.statusBar().showMessage(f"{ex} عذرا ,, يوجد خطأ في المدخلات ..يرجى التأكد من المعلومات")
        
        if mail_type == 14 :
            try:
                self.cur.execute('''UPDATE mail SET type_mail_id=%s,enter_date_hidden=%s,mail_date=%s,
                                mail_container=%s,mail_mostlm=%s,image=%s WHERE public_number_mail=%s''',[mail_type,eh,da,mad,mos,img,pn])
                self.db.commit()

                self.cur.execute('''UPDATE wared SET wared_from=%s,wared_to=%s WHERE wared_mail=%s''',[wfr,wto,pn])
                self.db.commit()

                self.cur.execute('''UPDATE personal_info SET name=%s,mobile=%s,national=%s,mahdar=%s,maksam=%s,manteka=%s,dareebe=%s,mehna=%s,segel=%s
                                        WHERE personal_info_mail=%s''',[name,mobile,national,mahder,maksam,manteka,dareebe,mehna,segel,pn])
                self.db.commit()
                self.statusBar().showMessage("تم تحديث بريد رواتب وأجور بنجاح")
            except Exception as ex:
                self.statusBar().showMessage(f"{ex} عذرا ,, يوجد خطأ في المدخلات ..يرجى التأكد من المعلومات")
        
        if mail_type == 15 or mail_type == 16 :
            try:
                self.cur.execute('''UPDATE mail SET type_mail_id=%s,enter_date_hidden=%s,mail_date=%s,
                                mail_container=%s,mail_mostlm=%s,image=%s WHERE public_number_mail=%s''',[mail_type,eh,da,mad,mos,img,pn])               
                self.db.commit()
                self.cur.execute('''UPDATE wared SET wared_from=%s,wared_to=%s WHERE wared_mail=%s''',[wfr,wto,pn])
                self.db.commit()
                self.cur.execute('''UPDATE personal_info SET name=%s,mobile=%s,national=%s,mahdar=%s,maksam=%s,manteka=%s,dareebe=%s,mehna=%s,segel=%s
                                        WHERE personal_info_mail=%s''',[name,mobile,national,mahder,maksam,manteka,dareebe,mehna,segel,pn])
                self.db.commit()
                self.statusBar().showMessage("تم تحديث بريد طلبات  بنجاح")
            except Exception as ex:
                self.statusBar().showMessage(f"{ex} عذرا ,, يوجد خطأ في المدخلات ..يرجى التأكد من المعلومات")
                # print(ex)

        if mail_type == 17 :
            try:
                self.cur.execute('''UPDATE mail SET type_mail_id=%s,enter_date_hidden=%s,mail_date=%s,
                                mail_container=%s,mail_mostlm=%s,image=%s WHERE public_number_mail=%s''',[mail_type,eh,da,mad,mos,img,pn])
                self.db.commit()
                self.cur.execute('''UPDATE sader SET sader_from=%s,sader_to=%s WHERE sader_mail=%s''',[sfr,sto,pn])
                self.db.commit()
                self.cur.execute('''UPDATE wared SET wared_from=%s,wared_to=%s WHERE wared_mail=%s''',[wfr,wto,pn])

                self.db.commit()
                self.statusBar().showMessage("تم تحديث بريد مراسلات الأقسام  بنجاح")
            except Exception as ex:
                self.statusBar().showMessage(f"{ex} عذرا ,, يوجد خطأ في المدخلات ..يرجى التأكد من المعلومات")
                # print(ex)
        
        #####
        updatemail = 6
        nameprocess = get_value_by_key(process_name_dict,updatemail)
        global employee_id
        #print(employee_id)
        datee = datetime.now()
        self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                         VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
        self.db.commit()
        self.Show_All_Operation(employee_id)
        self.Show_All_Mails()
        self.ClearDataSearch()
   
    ############//Employee DB//##############
    ##For Admin Only 
    
    def Show_All_Employee(self):
        self.tableWidget_2.clear()
        self.tableWidget_2.setColumnCount(7)
        column_name = ['الرقم الذاتي','الاسم','الموبايل','الايميل','كلمة السر','القسم','الدائرة']
        self.tableWidget_2.setHorizontalHeaderLabels(column_name)
        self.tableWidget_2.show()        
        self.cur.execute("""
                        SELECT id,name,mobile,email,password,emp_Dep,emp_Cir FROM employees
                        """)
        data = self.cur.fetchall()
        self.tableWidget_2.setRowCount(len(data))           
        for row , form in enumerate(data): 
            for col , item in enumerate(form):   
                if col == 5 :
                    value_dep = get_value_by_key(department_choices,item)
                    self.tableWidget_2.setItem(row,col,QTableWidgetItem(value_dep))
                else:  
                    self.tableWidget_2.setItem(row,col,QTableWidgetItem(str(item)))
                col+=1
            # row_pos = self.tableWidget_2.rowCount()
            # self.tableWidget_2.insertRow(row_pos)
    
    def Add_New_Employee(self):
        name = self.lineEdit_14.text()
        mobile = self.lineEdit_21.text().strip()
        

        email = self.lineEdit_5.text()
        
        if email != '':
            password = self.lineEdit_10.text()
            password2 = self.lineEdit_11.text()
            depart_emp = self.comboBox_4.currentText()
            id_depart_emp = get_key_by_value(department_choices,depart_emp)
            circle_emp = self.comboBox_5.currentText()
            if self.checkBox_5.isChecked():
                emp_per = 'abcdefghiklmn'
            else:
                emp_per = self.createpermision()
            try:
                if (password2 == password) and (name!= '')and (mobile!= '')and (password!= '')and (depart_emp!= '----'):
                    self.cur.execute("""
                    INSERT INTO employees (name,mobile,email,password,emp_Dep,emp_Cir,emp_Permisions) 
                    VALUES (%s,%s,%s,%s,%s,%s,%s)""",[name,mobile,email,password,id_depart_emp,circle_emp,emp_per])
                    self.db.commit()
                    self.statusBar().showMessage("تم إضافة موظف بنجاح")
                else:
                    self.statusBar().showMessage("  كلمتا السر غير متطابقتين أو أحد الحقول فارغة")
            except Exception as es:
                self.statusBar().showMessage(f"يوجد  خطأ في المعلومات المدخلة{es}")
            self.Show_All_Employee()
            self.Initilaize_Operation()
            self.clear_Cells()
            #####
            addemployee = 12
            nameprocess = get_value_by_key(process_name_dict,addemployee)
            global employee_id
            #print(employee_id)
            datee = datetime.now()
            self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                            VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
            self.db.commit()
        else :
            self.statusBar().showMessage("عنوان الايميل غير صالح")
            return
            
    def Search_Employee(self):

        try:
            index_combo = self.comboBox.currentIndex()
        except:
            self.statusBar().showMessage('الرجاء اختيار نوع البحث')

        if index_combo == 1 :#name
            name = self.lineEdit_15.text()
            try:
                self.cur.execute('''
                                SELECT * FROM employees WHERE name = %s
                                ''',[name])
                data = self.cur.fetchone()
                # print(data)
                empid = data[0]
                if data == None :
                    self.statusBar().showMessage("الموظف غير موجود في قاعدة البيانات")
            except Exception as ex:
                self.statusBar().showMessage(f"{ex}تأكد من المعلومات ثم أضغط بحث") 
                return
            
            
        elif index_combo == 2 :#mobile
            mobile_line = self.lineEdit_15
            if not self.validate_mobile_number(mobile_line):
                self.statusBar().showMessage("❌ رقم الهاتف غير صالح")
                return  # أوقف العملية بأمان
            mobile = mobile_line.text().strip()
            
            try:
                self.cur.execute('''
                                SELECT * FROM employees WHERE mobile = %s
                                ''',[mobile])
                data = self.cur.fetchone()
                empid = data[0]
                if data == None :
                    self.statusBar().showMessage("الموظف غير موجود في قاعدة البيانات")
            # print(data) 
            except Exception as ex:
                self.statusBar().showMessage(f"{ex}تأكد من المعلومات ثم أضغط بحث")
                return    
            
        elif index_combo == 3 :#email        
            email = self.lineEdit_15.text()
            # email = self.validate_email(self.lineEdit_15)
            if email !='':
                try:
                    self.cur.execute('''
                                    SELECT * FROM employees WHERE email = %s
                                    ''',[email])
                    data = self.cur.fetchone()
                    empid = data[0]
                    if data == None :
                        self.statusBar().showMessage("الموظف غير موجود في قاعدة البيانات")
                    # print(data)
                except Exception as ex:
                    self.statusBar().showMessage(f"{ex}تأكد من المعلومات ثم أضغط بحث") 
                    return
            else :
                self.statusBar().showMessage("يرجى مراجعة قسم المعلوماتية")

        if data:
            # print(data)
            self.lineEdit_61.setText(str(empid))
            self.lineEdit_29.setText(data[1])    
            self.lineEdit_36.setText(data[2])
            self.lineEdit_58.setText(data[3])
            # self.lineEdit_59.setText(data[4])
            self.lineEdit_60.setText(data[4])

            dep_name = get_value_by_key(department_choices,int(data[5]))
            # print(int(data[5]))
            # print(dep_name)
            self.comboBox_17.setCurrentText(dep_name)
            

            # id_cir = self.getIntCircleID(int(data[5]),data[6])
            
            self.comboBox_32.setCurrentText(data[6])
            emp_per = data[7]
            #print(emp_per)
            per = list(emp_per)
            #print(per)
            if len(per) == 13 :
                self.checkBox_5.setChecked(True)
            else :
                for c in per:
                    #employee
                    if c=='a':
                        self.checkBox_26.setChecked(True)
                    if c=='b':
                        self.checkBox.setChecked(True)
                    if c=='c':
                        self.checkBox_2.setChecked(True)
                    #circle
                    if c=='d':
                        self.checkBox_29.setChecked(True)
                    if c=='e':
                        self.checkBox_3.setChecked(True)
                    #bareed
                    if c=='f':
                        self.checkBox_34.setChecked(True)
                    if c=='g':
                        self.checkBox_35.setChecked(True)
                    if c=='h':
                        self.checkBox_36.setChecked(True) 
                    #combo
                    if c=='i':
                        self.checkBox_30.setChecked(True)
                    if c=='k':
                        self.checkBox_37.setChecked(True)
                    #setting
                    if c=='l':
                        self.checkBox_31.setChecked(True)   
                    if c=='m':
                        self.checkBox_32.setChecked(True)   
                    if c=='n':
                        self.checkBox_7.setChecked(True)   
                                    
    def Update_Employee(self):
        idemp = self.lineEdit_61.text()
        name = self.lineEdit_29.text()
        mobile_line = self.lineEdit_36
        if not self.validate_mobile_number(mobile_line):
            self.statusBar().showMessage("❌ رقم الهاتف غير صالح")
            return  # أوقف العملية بأمان
        mobile = mobile_line.text().strip()
        
        email = self.lineEdit_58.text()
        # email = self.validate_email(self.lineEdit_58)
        if email!= '':
            password2 = self.lineEdit_60.text()
            depart_emp = self.comboBox_17.currentText()
            depart_emp_id = get_key_by_value(department_choices,depart_emp)
            circle_emp = self.comboBox_32.currentText() 
            emp_per = self.createpermision()
                
            try:
                self.cur.execute("""
                UPDATE employees SET name=%s , mobile=%s , email=%s , password=%s , emp_Dep=%s , emp_Cir=%s , emp_Permisions=%s 
                WHERE id=%s""",[name,mobile,email,password2,depart_emp_id,circle_emp,emp_per,idemp])
                self.db.commit()
                updateemployee = 13
                nameprocess = get_value_by_key(process_name_dict,updateemployee)
                global employee_id
                datee = datetime.now()
                self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                self.db.commit()
                self.Show_All_Employee()
                self.clear_Cells()
                self.statusBar().showMessage("تم تحديث معلومات موظف بنجاح") 
            except:        
                self.statusBar().showMessage("تأكد من المعلومات ثم أضغط تحديث") 
                return 
        else :
                self.statusBar().showMessage("يرجى مراجعة قسم المعلوماتية")
                      
    def Del_Employee(self):
        try:    
            # row = self.tableWidget_2.currentRow()
            # column = self.tableWidget_2.currentColumn()
            # print(row)
            # print(column)
            # item = self.tableWidget_2.item( row, column)
            value_del = self.lineEdit_66.text()

            if value_del:
                reply = QMessageBox.critical(self, "تحذير", "هل تريد بالتأكيد الحذف?", QMessageBox.Yes | QMessageBox.No)
                if reply == QMessageBox.Yes:
                        self.cur.execute('''
                                    DELETE FROM employees WHERE id=%s
                                    ''',[value_del]) 
                        
                        #####
                        changeuserpass = 3
                        global employee_id
                        nameprocess = get_value_by_key(process_name_dict,changeuserpass)
                        #print(employee_id)
                        datee = datetime.now()
                        self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                        VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                        
                        
                        self.db.commit()
                        self.statusBar().showMessage("تم  حذف موظف بنجاح")
                else:
                    return
        except:        
            self.statusBar().showMessage("تأكد من المعلومات ثم أضغط بحث") 
            return
        
        self.Show_All_Employee()
    
    def clear_Cells(self):
        self.lineEdit_14.setText("")
        self.lineEdit_21.setText("")
        self.lineEdit_5.setText("")
        self.lineEdit_10.setText("")
        self.lineEdit_11.setText("")
        self.comboBox_4.setCurrentIndex(0)
        self.comboBox_5.setCurrentIndex(0)

        self.checkBox_26.setChecked(False)
        self.checkBox.setChecked(False)
        self.checkBox_7.setChecked(False)

        self.checkBox_29.setChecked(False)
        self.checkBox_3.setChecked(False)

        self.checkBox_34.setChecked(False)
        self.checkBox_35.setChecked(False)
        self.checkBox_36.setChecked(False)

        self.checkBox_30.setChecked(False)
        self.checkBox_37.setChecked(False)

        self.checkBox_31.setChecked(False)
        self.checkBox_32.setChecked(False)
        self.checkBox_7.setChecked(False)

        self.lineEdit_61.setText("")
        self.lineEdit_29.setText("")
        self.lineEdit_36.setText("")
        self.lineEdit_58.setText("")
        # self.lineEdit_59.setText("")
        self.lineEdit_60.setText("")

        self.comboBox_17.setCurrentIndex(0)
        self.comboBox_32.setCurrentIndex(0)

        self.lineEdit_15.setText('')

    def ChangeUserPass(self):
        try:
            name =  self.lineEdit_62.text()
            self.cur.execute('''SELECT password FROM employees WHERE name=%s''',[name])
            password= self.cur.fetchone()
            #print(password)
            if password[0] == self.lineEdit_64.text() :
                pass1 = self.lineEdit_3.text()
                pass2 = self.lineEdit_4.text()
                if pass1 == pass2:
                    self.cur.execute('''UPDATE employees SET password=%s WHERE name=%s''',[pass1,name])
                    
                    #####
                    updateemployee = 14
                    nameprocess = get_value_by_key(process_name_dict,updateemployee)
                    global employee_id
                    #print(employee_id)
                    datee = datetime.now()
                    self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                                    VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
                    
                    
                    self.db.commit()
                    self.statusBar().showMessage("تم تغيير كلمة السر بنجاح")
                else:
                    self.statusBar().showMessage("كلمتا السر غير متطابقتين") 
            else:
                self.statusBar().showMessage("تأكد من كلمة السر ثم إعد المحاولة") 
        except Exception as e:
            self.statusBar().showMessage("تأكد من معلوماتك قبل الضغط ")

    def ChangeAdminPass(self):
        self.cur.execute('''SELECT password FROM employees WHERE name="admin"''')
        password= self.cur.fetchone()
        # print(password)
        if password[0] == self.lineEdit_63.text() :
            pass1 = self.lineEdit_12.text()
            pass2 = self.lineEdit_13.text()
            if pass1 == pass2:
                self.cur.execute('''UPDATE employees SET password=%s WHERE name="admin"''',[pass1])
                self.db.commit()
                self.statusBar().showMessage("تم تغيير كلمة السر بنجاح")
            else:
                self.statusBar().showMessage("كلمتا السر غير متطابقتين") 
        else:
            self.statusBar().showMessage("تأكد من كلمة السر الادمن ثم إعد المحاولة") 
  
##################################
            ###Image
    def BrowseImage(self):
        file_path, _ = QFileDialog.getOpenFileName(self,"اختر صورة أو مستند","",
                                                   "Images and Text Files (*.png *.jpg *.jpeg *.bmp *.gif *.txt *.doc *.docx)")
        if file_path:
            self.file_img_txt_path = file_path
            #print(file_path)
            self.statusBar().showMessage(f" مسار الصورة: {file_path}")

        
            extension = os.path.splitext(file_path)[1].lower()

            # إخفاء المعاينات مؤقتًا
            self.text_area.setVisible(False)
            self.image_label.setVisible(False)

            if extension in ['.txt']:
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        #self.text_area.setProperity
                        self.text_area.setText(content)
                        self.text_area.setVisible(True)
                except Exception as e:
                    self.text_area.setText(f"خطأ في قراءة الملف:\n{e}")
                    self.text_area.setVisible(True)

            elif extension in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
                pixmap = QPixmap(file_path)
                self.image_label.setPixmap(pixmap)
                self.image_label.setVisible(True)

            else:
                self.text_area.setText("نوع الملف غير مدعوم للمعاينة.")
                self.text_area.setVisible(True)

    def open_capture_window(self):
        self.capture_window = CaptureWindow()
        self.capture_window.path_ready.connect(self.handle_image_data)
       
        self.capture_window.path_ready.connect(self.show_image_in_label) # ربط الإشارة
        self.capture_window.show()

    def handle_image_data(self, image_path, image_bytes):
        # print("📁 المسار:", image_path)
        self.statusBar().showMessage(f"📁 المسار:{image_path}🧬 حجم الصورة:{len(image_bytes)}") 

        # print("🧬 حجم الصورة:", len(image_bytes))
        self.current_img_byte = image_bytes
        self.current_image_path = image_path
        # QMessageBox.information(self, "📸 تم الالتقاط", "تم تخزين الصورة مؤقتًا، اضغط حفظ لإرسالها.")

    def show_image_in_label2(self, path):
        if os.path.exists(path) and path.lower().endswith(('.jpg', '.jpeg', '.png')):
            pixmap = QPixmap(path)
            scaled = pixmap.scaled(self.image_label_2.size(), Qt.KeepAspectRatio, Qt.SmoothTransformation)
            self.image_label_2.setPixmap(scaled)
        else:
            pass
            #print(f"❌ الملف غير موجود أو ليس صورة: {path}")

    def show_image_in_label(self, path):
        pixmap = QPixmap(path)
        scaled = pixmap.scaled(self.label_preview.size(), Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.label_preview.setPixmap(scaled)
        self.current_image_path = path  # نخزن المسار لاستخدامه عند تغيير الحجم

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if self.current_image_path:
            self.show_image_in_label(self.current_image_path)
 
#############//Emplyee Report//#############
# For Admin And Employee
#############//Monthly Report//###############

    def FilterMails(self):
        index = self.comboBox_8.currentIndex()
        
        if index ==1 or index ==2 or index ==3 or index ==4 or index ==5 or index == 6:
            self.tableWidget_4.clear()
            self.tableWidget_4.setColumnCount(16)
            column_name = ['الرقم العام','تاريخ البريد','مضمون البريد','مستلم البريد','نسخة/صورة','صادر من','صادر إلى',
                           'الاسم الشخصي','الموبايل','الرقم الوطني','المحضر','المقسم','المنطقة','الرقم الضريبي','المهنة','السجل' ]
            self.tableWidget_4.setHorizontalHeaderLabels(column_name)
            self.tableWidget_4.show()
            # self.tableWidget_4.insertRow(0)
            
            self.cur.execute('''SELECT 
                                    mail.public_number_mail,
                                    mail.mail_date,
                                    mail.mail_container,
                                    mail.mail_mostlm,
                                    mail.image,
                                    sader.sader_from,
                                    sader.sader_to,
                                    personal_info.name,
                                    personal_info.mobile,
                                    personal_info.national,
                                    personal_info.mahdar,
                                    personal_info.maksam,
                                    personal_info.manteka,
                                    personal_info.dareebe,
                                    personal_info.mehna,
                                    personal_info.segel
                                FROM mail
                                LEFT JOIN sader ON mail.public_number_mail = sader.sader_mail
                                LEFT JOIN personal_info ON mail.public_number_mail = personal_info.personal_info_mail
                                WHERE mail.type_mail_id = %s;
                            ''', [index])

            mail_data = self.cur.fetchall()
            self.tableWidget_4.setRowCount(len(mail_data))           

            # print(mail_data)
            for row , form in enumerate(mail_data):
                    for col , item in enumerate(form):
                        

                        self.tableWidget_4.setItem(row,col,QTableWidgetItem(str(item)))
                        

                        col+=1
                    # row_pos = self.tableWidget_4.rowCount()
                    # self.tableWidget_4.insertRow(row_pos)      
        
        if index == 7 or index == 8 or index == 9 or index == 10:
            self.tableWidget_4.clear()
            self.tableWidget_4.setColumnCount(16)
            column_name = ['الرقم العام','تاريخ البريد','مضمون البريد','مستلم البريد','نسخة/صورة','وارد من','وارد إلى',
                           'الاسم الشخصي','الموبايل','الرقم الوطني','المحضر','المقسم','المنطقة','الرقم الضريبي','المهنة','السجل']
            self.tableWidget_4.setHorizontalHeaderLabels(column_name)
            self.tableWidget_4.show()
            # self.tableWidget_4.insertRow(0)
            #print(mail_type_tuble[0][1])
            self.cur.execute('''SELECT 
                                    mail.public_number_mail,
                                    mail.mail_date,
                                    mail.mail_container,
                                    mail.mail_mostlm,
                                    mail.image,
                                    wared.wared_from,
                                    wared.wared_to,
                                    personal_info.name,
                                    personal_info.mobile,
                                    personal_info.national,
                                    personal_info.mahdar,
                                    personal_info.maksam,
                                    personal_info.manteka,
                                    personal_info.dareebe,
                                    personal_info.mehna,
                                    personal_info.segel
                                FROM mail
                                LEFT JOIN wared ON mail.public_number_mail = wared.wared_mail
                                LEFT JOIN personal_info ON mail.public_number_mail = personal_info.personal_info_mail
                                WHERE mail.type_mail_id = %s;
                            ''', [index])
            mail_data = self.cur.fetchall()
            # print(mail_data)
            self.tableWidget_4.setRowCount(len(mail_data))           

            for row , form in enumerate(mail_data):
                    for col , item in enumerate(form):
                        self.tableWidget_4.setItem(row,col,QTableWidgetItem(str(item)))
                        col+=1

                    # row_pos = self.tableWidget_4.rowCount()
                    # self.tableWidget_4.insertRow(row_pos)      
        
        if index == 11:
            self.tableWidget_4.clear()
            self.tableWidget_4.setColumnCount(21)
            column_name = ['الرقم العام','تاريخ البريد','مضمون البريد','مستلم البريد','نسخة/صورة',
                           'نوع الحكم','صاحب الدعوى','اسم المدعى عليه','رقم القرار','رقم الاساس',
                           'رقم التنفيذ','تاريخ الحكم',
                           'الاسم الشخصي','الموبايل','الرقم الوطني','المحضر','المقسم','المنطقة','الرقم الضريبي','المهنة','السجل']
            self.tableWidget_4.setHorizontalHeaderLabels(column_name)
            self.tableWidget_4.show()
            # self.tableWidget_4.insertRow(0)
            #print(mail_type_tuble[0][1])
            self.cur.execute('''SELECT mail.public_number_mail,mail.mail_date,mail.mail_container,mail.mail_mostlm,mail.image,
                             ahkam.ahkam_type,ahkam.mh_name,ahkam.mh3_name,ahkam.karar,ahkam.asas,ahkam.tanfeez,ahkam.date_ahkam,
                             personal_info.name,personal_info.mobile,personal_info.national,personal_info.mahdar,personal_info.maksam,
                             personal_info.manteka,personal_info.dareebe,personal_info.mehna,personal_info.segel FROM mail
                             LEFT JOIN ahkam ON mail.public_number_mail = ahkam.ahkam_mail
                             LEFT JOIN personal_info ON mail.public_number_mail = personal_info.personal_info_mail
                                WHERE mail.type_mail_id = %s;''', [index])
            mail_data = self.cur.fetchall()
            # print(mail_data)
            self.tableWidget_4.setRowCount(len(mail_data))           

            for row , form in enumerate(mail_data):
                    for col , item in enumerate(form):
                        if col == 6 :
                            self.cur.execute('''SELECT name FROM cb_ahkam WHERE id=%s''',[int(item)])
                            name = self.cur.fetchone()
                            self.tableWidget_4.setItem(row,col,QTableWidgetItem(name[0]))
                        else:
                            self.tableWidget_4.setItem(row,col,QTableWidgetItem(str(item)))
                        col+=1

                    # row_pos = self.tableWidget_4.rowCount()
                    # self.tableWidget_4.insertRow(row_pos)      
        
        if index == 12:
            self.tableWidget_4.clear()
            self.tableWidget_4.setColumnCount(16)
            column_name = ['الرقم العام','تاريخ البريد','مضمون البريد','مستلم البريد','نسخة/صورة','صادر من','صادر إلى',
                           'الاسم الشخصي','الموبايل','الرقم الوطني','المحضر','المقسم','المنطقة','الرقم الضريبي','المهنة','السجل']
            self.tableWidget_4.setHorizontalHeaderLabels(column_name)
            self.tableWidget_4.show()
            # self.tableWidget_4.insertRow(0)
            #print(mail_type_tuble[0][1])
            self.cur.execute('''SELECT 
                                    mail.public_number_mail,
                                    mail.mail_date,
                                    mail.mail_container,
                                    mail.mail_mostlm,
                                    mail.image,
                                    sader.sader_from,
                                    sader.sader_to,
                                    personal_info.name,
                                    personal_info.mobile,
                                    personal_info.national,
                                    personal_info.mahdar,
                                    personal_info.maksam,
                                    personal_info.manteka,
                                    personal_info.dareebe,
                                    personal_info.mehna,
                                    personal_info.segel
                                FROM mail
                                LEFT JOIN sader ON mail.public_number_mail = sader.sader_mail
                                LEFT JOIN personal_info ON mail.public_number_mail = personal_info.personal_info_mail
                                WHERE mail.type_mail_id = %s;
                            ''', [index])
            mail_data = self.cur.fetchall()
            #print(mail_data)
            self.tableWidget_4.setRowCount(len(mail_data)) 
            for row , form in enumerate(mail_data):
                    for col , item in enumerate(form):
                        self.tableWidget_4.setItem(row,col,QTableWidgetItem(str(item)))
                        col+=1

                    # row_pos = self.tableWidget_4.rowCount()
                    # self.tableWidget_4.insertRow(row_pos)      
       
        if index == 13:
            self.tableWidget_4.clear()
            self.tableWidget_4.setColumnCount(19)
            column_name = ['الرقم العام','تاريخ البريد','مضمون البريد','مستلم البريد','نسخة/صورة','وارد من','وارد إلى',
                           'الاسم الشخصي','الموبايل','الرقم الوطني','المحضر','المقسم','المنطقة','الرقم الضريبي','المهنة','السجل'
                           'اسم المعترض','تاريخ الاعتراض',' معترض على '
                           ]
            self.tableWidget_4.setHorizontalHeaderLabels(column_name)
            self.tableWidget_4.show()
            # self.tableWidget_4.insertRow(0)
            #print(mail_type_tuble[0][1])
            self.cur.execute('''SELECT 
                                    mail.public_number_mail,
                                    mail.mail_date,
                                    mail.mail_container,
                                    mail.mail_mostlm,
                                    mail.image,
                                    wared.wared_from,
                                    wared.wared_to,
                                    mo3tred.mo3tred_name,
                                    mo3tred.mo3tred_date,
                                    mo3tred.mo3tred_at,
                                    personal_info.name,
                                    personal_info.mobile,
                                    personal_info.national,
                                    personal_info.mahdar,
                                    personal_info.maksam,
                                    personal_info.manteka,
                                    personal_info.dareebe,
                                    personal_info.mehna,
                                    personal_info.segel
                                FROM mail
                                LEFT JOIN wared ON mail.public_number_mail = wared.wared_mail
                                LEFT JOIN mo3tred ON mail.public_number_mail = mo3tred.mo3tred_mail 
                                LEFT JOIN personal_info ON mo3tred.mo3tred_per = personal_info.idpersonal_info
                                WHERE mail.type_mail_id = %s;
                            ''', [index])

            

            
            mail_data = self.cur.fetchall()
            # print(mail_data)
            self.tableWidget_4.setRowCount(len(mail_data)) 
            for row , form in enumerate(mail_data):
                    for col , item in enumerate(form):
                        self.tableWidget_4.setItem(row,col,QTableWidgetItem(str(item)))
                        col+=1

                    # row_pos = self.tableWidget_4.rowCount()
                    # self.tableWidget_4.insertRow(row_pos)      
            
        if index == 14:
            self.tableWidget_4.clear()
            self.tableWidget_4.setColumnCount(16)
            column_name = ['الرقم العام','تاريخ البريد','مضمون البريد','مستلم البريد','نسخة/صورة','وارد من','وارد إلى',
                           'الاسم الشخصي','الموبايل','الرقم الوطني','المحضر','المقسم','المنطقة','الرقم الضريبي','المهنة','السجل']
            self.tableWidget_4.setHorizontalHeaderLabels(column_name)
            self.tableWidget_4.show()
            # self.tableWidget_4.insertRow(0)
            #print(mail_type_tuble[0][1])
            self.cur.execute('''SELECT 
                                    mail.public_number_mail,
                                    mail.mail_date,
                                    mail.mail_container,
                                    mail.mail_mostlm,
                                    mail.image,
                                    wared.wared_from,
                                    wared.wared_to,
                                    personal_info.name,
                                    personal_info.mobile,
                                    personal_info.national,
                                    personal_info.mahdar,
                                    personal_info.maksam,
                                    personal_info.manteka,
                                    personal_info.dareebe,
                                    personal_info.mehna,
                                    personal_info.segel
                                FROM mail
                                LEFT JOIN wared ON mail.public_number_mail = wared.wared_mail
                                LEFT JOIN personal_info ON mail.public_number_mail = personal_info.personal_info_mail
                                WHERE mail.type_mail_id = %s;
                            ''', [index])
            mail_data = self.cur.fetchall()
            # print(mail_data)
            self.tableWidget_4.setRowCount(len(mail_data)) 
            for row , form in enumerate(mail_data):
                    for col , item in enumerate(form):
                        self.tableWidget_4.setItem(row,col,QTableWidgetItem(str(item)))
                        col+=1

                    row_pos = self.tableWidget_4.rowCount()
                    self.tableWidget_4.insertRow(row_pos)      
    
        if index == 15 or index == 16:
            self.tableWidget_4.clear()
            self.tableWidget_4.setColumnCount(16)
            column_name = ['الرقم العام','تاريخ البريد','مضمون البريد','مستلم البريد','نسخة/صورة','وارد من','وارد إلى',
                           'الاسم الشخصي','الموبايل','الرقم الوطني','المحضر','المقسم','المنطقة','الرقم الضريبي','المهنة','السجل']
            self.tableWidget_4.setHorizontalHeaderLabels(column_name)
            self.tableWidget_4.show()
            # self.tableWidget_4.insertRow(0)
            #print(mail_type_tuble[0][1])
            self.cur.execute('''SELECT 
                                    mail.public_number_mail,
                                    mail.mail_date,
                                    mail.mail_container,
                                    mail.mail_mostlm,
                                    mail.image,
                                    wared.wared_from,
                                    wared.wared_to,
                                    personal_info.name,
                                    personal_info.mobile,
                                    personal_info.national,
                                    personal_info.mahdar,
                                    personal_info.maksam,
                                    personal_info.manteka,
                                    personal_info.dareebe,
                                    personal_info.mehna,
                                    personal_info.segel
                                FROM mail
                                LEFT JOIN wared ON mail.public_number_mail = wared.wared_mail
                                LEFT JOIN personal_info ON mail.public_number_mail = personal_info.personal_info_mail
                                WHERE mail.type_mail_id = %s;
                            ''', [index])
            mail_data = self.cur.fetchall()
            # print(mail_data)
            self.tableWidget_4.setRowCount(len(mail_data)) 
            for row , form in enumerate(mail_data):
                    for col , item in enumerate(form):
                        self.tableWidget_4.setItem(row,col,QTableWidgetItem(str(item)))
                        col+=1

                    # row_pos = self.tableWidget_4.rowCount()
                    # self.tableWidget_4.insertRow(row_pos)      
     
        if index == 17:
            self.tableWidget_4.clear()
            self.tableWidget_4.setColumnCount(7)
            column_name = ['الرقم العام','تاريخ البريد','مضمون البريد','مستلم البريد','نسخة/صورة','كيواركود','من قسم','إلى قسم']
            self.tableWidget_4.setHorizontalHeaderLabels(column_name)
            self.tableWidget_4.show()
            # self.tableWidget_4.insertRow(0)
            #print(mail_type_tuble[0][1])
            self.cur.execute('''
                        SELECT mail.public_number_mail , mail.mail_date , mail.mail_container , mail.mail_mostlm ,
                               mail.image ,aksam.from_k,aksam.to_k FROM mail 
                             JOIN aksam ON mail.public_number_mail = aksam.aksam_mail
                             WHERE mail.type_mail_id = %s; ''',[index])
            mail_data = self.cur.fetchall()
            # print(mail_data)
            self.tableWidget_4.setRowCount(len(mail_data)) 
            for row , form in enumerate(mail_data):
                    for col , item in enumerate(form):
                        self.tableWidget_4.setItem(row,col,QTableWidgetItem(str(item)))
                        col+=1

                    # row_pos = self.tableWidget_4.rowCount()
                    # self.tableWidget_4.insertRow(row_pos)      

        if index == 18:
            self.tableWidget_4.clear()
            self.tableWidget_4.setColumnCount(5)
            column_name = ['الرقم العام','تاريخ البريد','مضمون البريد','مستلم البريد','نسخة/صورة']
            self.tableWidget_4.setHorizontalHeaderLabels(column_name)
            self.tableWidget_4.show()
            self.tableWidget_4.insertRow(0)
            #print(mail_type_tuble[0][1])
            self.cur.execute('''
                        SELECT public_number_mail , mail_date , mail_container , mail_mostlm ,
                               image  FROM mail
                             WHERE type_mail_id = %s; ''',[index])
            mail_data = self.cur.fetchall()
            # print(mail_data)
            for row , form in enumerate(mail_data):
                    for col , item in enumerate(form):
                        self.tableWidget_4.setItem(row,col,QTableWidgetItem(str(item)))
                        col+=1

                    row_pos = self.tableWidget_4.rowCount()
                    self.tableWidget_4.insertRow(row_pos)      

        
         ###  
        #####
        filters1 = 7
        nameprocess = get_value_by_key(process_name_dict,filters1)
        global employee_id
        #print(employee_id)
        datee = datetime.now()
        self.cur.execute('''INSERT INTO dailymovements (process_type,emp_id,date_process )
                        VALUES (%s,%s,%s)''',[nameprocess,employee_id,datee])
        self.db.commit()
                  
    def Export_cvx(self):
        # فتح نافذة لحفظ الملف
        file_path, _ = QFileDialog.getSaveFileName(
            self, "اختر مكان الحفظ", "", "CSV Files (*.csv)"
        )
        if not file_path:
            return

        row_count = self.tableWidget_4.rowCount()
        column_count = self.tableWidget_4.columnCount()

        headers = [
            self.tableWidget_4.horizontalHeaderItem(col).text()
            if self.tableWidget_4.horizontalHeaderItem(col)
            else f"عمود_{col + 1}"
            for col in range(column_count)
        ]

        data = []
        for row in range(row_count):
            row_data = [
                self.tableWidget_4.item(row, col).text() if self.tableWidget_4.item(row, col) else ''
                for col in range(column_count)
            ]
            data.append(row_data)

        # إنشاء ملف بصيغة CSV مع تنسيق رمزي
        with open(file_path, 'w', encoding='utf-8-sig', newline='') as f:
            # إدراج نص رمزي يمثل الشعار السوري في بداية الملف
            f.write("🦅 الجمهورية العربية السورية - شعار الدولة\n\n")
            
            # كتابة الجدول
            df = pd.DataFrame(data, columns=headers)
            df.to_csv(f, index=False)

        self.statusBar().showMessage(f"✅ تم حفظ الملف بصيغة CSV مع تنسيق رمزي{file_path}")

    def Export_xlsx2(self):
        # فتح نافذة لحفظ الملف
        file_path, _ = QFileDialog.getSaveFileName(
            self, "اختر مكان الحفظ", "", "Excel Files (*.xlsx)"
        )
        if not file_path:
            return

        try:
            # إعداد ملف الإكسل
            workbook = xlsxwriter.Workbook(file_path)
            worksheet = workbook.add_worksheet("البيانات")

            # إدراج الصورة
            image_path = os.path.join(os.getcwd(), "syrianar3kab.png")
            worksheet.insert_image('A1', image_path, {
                'x_scale': 0.7,
                'y_scale': 0.7,
                'object_position': 1  # يسمح بتحريك الصورة مع الخلايا
            })

            # تحديد بداية الجدول بعد الصورة
            start_row = 15  # يمكن تغييره حسب حجم الصورة

            column_count = self.tableWidget_4.columnCount()
            row_count = self.tableWidget_4.rowCount()
            # print("عدد الصفوف:", row_count)
            # print("عدد الأعمدة:", column_count)
            if row_count == 0 or column_count == 0:
                self.statusBar().showMessage("⚠️ لا توجد بيانات في الجدول للتصدير.")
                return


            # تنسيق العناوين
            header_format = workbook.add_format({'bold': True, 'bg_color': '#D7E4BC', 'align': 'center', 'border': 1})
            cell_format = workbook.add_format({'align': 'center', 'border': 1})

            # كتابة رؤوس الأعمدة
            for col in range(column_count):
                header_text = self.tableWidget_4.horizontalHeaderItem(col).text() if self.tableWidget_4.horizontalHeaderItem(col) else f"عمود_{col + 1}"
                worksheet.write(start_row, col, header_text, header_format)

            # كتابة البيانات
            for row in range(row_count):
                for col in range(column_count):
                    item = self.tableWidget_4.item(row, col)
                    # print(f"({row},{col}) =", item.text() if item else "None")
                    value = item.text() if item else ''
                    worksheet.write(start_row + 1 + row, col, value, cell_format)

            # ضبط عرض الأعمدة تلقائيًا
            for col in range(column_count):
                max_length = len(self.tableWidget_4.horizontalHeaderItem(col).text()) if self.tableWidget_4.horizontalHeaderItem(col) else 10
                for row in range(row_count):
                    item = self.tableWidget_4.item(row, col)
                    if item and item.text():
                        max_length = max(max_length, len(item.text()))
                worksheet.set_column(col, col, max_length + 2)

            # حفظ الملف
            workbook.close()
            self.statusBar().showMessage("✅ تم حفظ الملف مع الصورة والبيانات: " + file_path)

        except Exception as e:
            self.statusBar().showMessage(f"⚠️ فشل التصدير: {e}")
            self.statusBar().showMessage(f"⚠️ فشل التصدير: {e}")
      
##############################################
################//creating ui object//########

def main():
    app = QApplication(sys.argv)
    window = Main()
    window.show()
    app.exec_()

#############################################
################//executing program//########    

if __name__ == '__main__':
    try:
        elevate(show_console=True)
        main()
    except Exception:
        with open("error_log.txt", "w", encoding="utf-8") as f:
            f.write(traceback.format_exc())
        sys.exit(1)







    



    


